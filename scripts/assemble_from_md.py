#!/usr/bin/env python3
"""
Convert annotated Markdown to DOCX in-place.

Parses annotation blocks:
  - <!-- <updated section="X" para="N"> --> ... <!-- </updated> -->
  - <!-- <to-be-removed section="X" para="N" reason="..."> --> ... <!-- </to-be-removed> -->
  - <!-- <human-task type="diagram" section="X" ...> --> ... <!-- </human-task> -->
  - <!-- <fact-check status="misaligned" ...> --> ... <!-- </fact-check> -->

Applies updates/removals to a single .docx file in-place.
Emits JSONL log to .manuscript-revision.log alongside the .docx.
"""

import argparse
import json
import re
import sys
from dataclasses import dataclass, asdict, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


@dataclass
class LogEntry:
    ts: str
    section: str
    para: Optional[int] = None
    figure: Optional[str] = None
    action: str = ""
    type: Optional[str] = None
    reason: Optional[str] = None
    status: str = "ok"
    detail: Optional[str] = None

    def to_json(self) -> str:
        # Remove None values for cleaner JSONL
        d = {k: v for k, v in asdict(self).items() if v is not None}
        return json.dumps(d, ensure_ascii=False)


@dataclass
class UpdatedBlock:
    section: str
    para: int
    content: str
    raw: str  # Full block text for potential cleanup


@dataclass
class ToBeRemovedBlock:
    section: str
    para: int
    reason: str
    content: str
    raw: str


@dataclass
class HumanTaskBlock:
    type: str
    section: str
    figure: Optional[str] = None
    table: Optional[str] = None
    description: str = ""
    raw: str = ""


@dataclass
class FactCheckBlock:
    status: str
    claim: str
    bank_section: Optional[str] = None
    bank_argument: Optional[str] = None
    detail: str = ""
    source: Optional[str] = None
    raw: str = ""


def parse_annotation_blocks(markdown: str):
    """Parse all annotation blocks from markdown. Returns lists of each block type."""
    updated_blocks = []
    removed_blocks = []
    human_task_blocks = []
    fact_check_blocks = []

    # Pattern for <!-- <tag attr="val" ...> --> ... <!-- </tag> -->
    # Matches opening tag, content, closing tag
    # Tag names can contain hyphens (e.g., to-be-removed, fact-check, human-task)
    block_pattern = re.compile(
        r'<!--\s*<([\w-]+)(?:\s+([^>]+?))?>\s*-->'
        r'(.*?)'
        r'<!--\s*</\1>\s*-->',
        re.DOTALL
    )

    for match in block_pattern.finditer(markdown):
        tag = match.group(1).lower()
        attrs_str = match.group(2) or ""
        content = match.group(3).strip()
        raw = match.group(0)

        # Parse attributes
        attrs = {}
        attr_pattern = re.compile(r'(\w+)="([^"]*)"')
        for attr_match in attr_pattern.finditer(attrs_str):
            attrs[attr_match.group(1)] = attr_match.group(2)

        if tag == "updated":
            section = attrs.get("section", "").strip()
            para_str = attrs.get("para", "").strip()
            if not section or not para_str:
                print(f"⚠️  Invalid <updated> block: missing section or para")
                continue
            try:
                para = int(para_str)
            except ValueError:
                print(f"⚠️  Invalid <updated> block: para must be integer")
                continue
            updated_blocks.append(UpdatedBlock(
                section=section, para=para, content=content, raw=raw
            ))

        elif tag == "to-be-removed":
            section = attrs.get("section", "").strip()
            para_str = attrs.get("para", "").strip()
            reason = attrs.get("reason", "").strip()
            if not section or not para_str or not reason:
                print(f"⚠️  Invalid <to-be-removed> block: missing section, para, or reason")
                continue
            try:
                para = int(para_str)
            except ValueError:
                print(f"⚠️  Invalid <to-be-removed> block: para must be integer")
                continue
            removed_blocks.append(ToBeRemovedBlock(
                section=section, para=para, reason=reason, content=content, raw=raw
            ))

        elif tag == "human-task":
            task_type = attrs.get("type", "").strip()
            section = attrs.get("section", "").strip()
            figure = attrs.get("figure", "").strip() or None
            table = attrs.get("table", "").strip() or None
            description = attrs.get("description", "").strip()
            if not task_type or not section:
                print(f"⚠️  Invalid <human-task> block: missing type or section")
                continue
            human_task_blocks.append(HumanTaskBlock(
                type=task_type, section=section, figure=figure,
                table=table, description=description, raw=raw
            ))

        elif tag == "fact-check":
            status = attrs.get("status", "").strip()
            claim = attrs.get("claim", "").strip()
            bank_section = attrs.get("bank_section", "").strip() or None
            bank_argument = attrs.get("bank_argument", "").strip() or None
            detail = attrs.get("detail", "").strip()
            source = attrs.get("source", "").strip() or None
            if not status or not claim:
                print(f"⚠️  Invalid <fact-check> block: missing status or claim")
                continue
            fact_check_blocks.append(FactCheckBlock(
                status=status, claim=claim, bank_section=bank_section,
                bank_argument=bank_argument, detail=detail, source=source, raw=raw
            ))

    return updated_blocks, removed_blocks, human_task_blocks, fact_check_blocks


def is_heading(para: Paragraph, known_headings: set[str]) -> bool:
    """Detect if a paragraph is a section heading."""
    text = para.text.strip()
    # Check by style
    if para.style.name.startswith("Heading"):
        return True
    # Check by known heading texts (for docs without proper styles)
    if text in known_headings:
        return True
    # Check by common patterns
    if text.startswith("Chapter ") and text[8:].strip().isdigit():
        return True
    if text in ("INTRODUCTION", "METHODOLOGY", "REFERENCES", "APPENDICES"):
        return True
    if text.startswith("Figure ") or text.startswith("Table "):
        return False  # These are figure/table captions, not section headings
    # Heuristic: short line, title case or upper case, no period at end
    if len(text) < 80 and text and not text.endswith("."):
        words = text.split()
        if 1 < len(words) <= 8:
            # Likely a heading if mostly title/upper case
            title_case = sum(1 for w in words if w[0].isupper() and w[1:].islower())
            upper_case = sum(1 for w in words if w.isupper())
            if title_case + upper_case >= len(words) * 0.7:
                return True
    return False


def find_section_paragraphs(doc: Document, section_heading: str):
    """
    Find all body paragraphs belonging to a section.
    Returns list of (paragraph_index, paragraph_object) for paragraphs
    under the given section heading (exact match, case-sensitive).
    Works with both styled headings and text-based headings.
    """
    # Build set of all heading-like texts in document for detection
    known_headings = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not text.endswith(".") and len(text) < 100:
            known_headings.add(text)

    in_section = False
    section_paragraphs = []
    para_idx = 0

    for para in doc.paragraphs:
        # Check if this is a heading
        if is_heading(para, known_headings):
            heading_text = para.text.strip()
            if heading_text == section_heading:
                in_section = True
                para_idx = 0
            elif in_section:
                # Hit next heading - section ended
                break
            continue

        if in_section:
            # Skip empty paragraphs
            if para.text.strip():
                para_idx += 1
                section_paragraphs.append((para_idx, para))

    return section_paragraphs


def fuzzy_match(text1: str, text2: str, threshold: float = 0.8) -> bool:
    """Simple fuzzy match using word overlap ratio."""
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())
    if not words1 or not words2:
        return False
    intersection = words1 & words2
    union = words1 | words2
    return len(intersection) / len(union) >= threshold


def apply_updated_blocks(doc: Document, blocks: list[UpdatedBlock], log_path: Path):
    """Apply <updated> blocks to document. Returns list of log entries."""
    logs = []

    for block in blocks:
        section_paragraphs = find_section_paragraphs(doc, block.section)

        if not section_paragraphs:
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="updated",
                status="failed",
                detail=f"Section '{block.section}' not found"
            ))
            continue

        if block.para > len(section_paragraphs):
            # Try fallback: fuzzy match nearby paragraphs
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="updated",
                status="failed",
                detail=f"Paragraph {block.para} out of range (section has {len(section_paragraphs)} paragraphs)"
            ))
            continue

        # Get target paragraph (1-indexed)
        _, target_para = section_paragraphs[block.para - 1]

        try:
            # Replace text while preserving runs/style as much as possible
            # Clear existing runs and add new one with the content
            for run in target_para.runs:
                run.text = ""
            if target_para.runs:
                target_para.runs[0].text = block.content
            else:
                target_para.add_run(block.content)

            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="updated",
                status="ok"
            ))
        except Exception as e:
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="updated",
                status="failed",
                detail=str(e)
            ))

    return logs


def apply_removed_blocks(doc: Document, blocks: list[ToBeRemovedBlock], log_path: Path):
    """Apply <to-be-removed> blocks to document. Returns list of log entries."""
    logs = []

    # We need to track paragraph removals carefully since indices shift
    # Collect all removals first, then apply from bottom to top
    removals = []  # (section, para_idx_in_section, paragraph_object, reason)

    for block in blocks:
        section_paragraphs = find_section_paragraphs(doc, block.section)

        if not section_paragraphs:
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="removed",
                reason=block.reason,
                status="failed",
                detail=f"Section '{block.section}' not found"
            ))
            continue

        if block.para > len(section_paragraphs):
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=block.section,
                para=block.para,
                action="removed",
                reason=block.reason,
                status="failed",
                detail=f"Paragraph {block.para} out of range (section has {len(section_paragraphs)} paragraphs)"
            ))
            continue

        _, target_para = section_paragraphs[block.para - 1]
        removals.append((block.section, block.para - 1, target_para, block.reason))

    # Apply removals from last to first (to preserve indices)
    for section, para_idx, target_para, reason in reversed(removals):
        try:
            # Remove paragraph from document XML
            p_element = target_para._element
            p_element.getparent().remove(p_element)

            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=section,
                para=para_idx + 1,
                action="removed",
                reason=reason,
                status="ok"
            ))
        except Exception as e:
            logs.append(LogEntry(
                ts=datetime.utcnow().isoformat() + "Z",
                section=section,
                para=para_idx + 1,
                action="removed",
                reason=reason,
                status="failed",
                detail=str(e)
            ))

    return logs


def log_human_tasks(blocks: list[HumanTaskBlock]):
    """Log <human-task> blocks as pending. Returns list of log entries."""
    logs = []
    for block in blocks:
        logs.append(LogEntry(
            ts=datetime.utcnow().isoformat() + "Z",
            section=block.section,
            figure=block.figure,
            action="human-task",
            type=block.type,
            status="pending",
            detail=block.description
        ))
    return logs


def log_fact_checks(blocks: list[FactCheckBlock]):
    """Log <fact-check> blocks for review. Returns list of log entries."""
    logs = []
    for block in blocks:
        logs.append(LogEntry(
            ts=datetime.utcnow().isoformat() + "Z",
            section=block.bank_section or "unknown",
            action="fact-check",
            status=block.status,
            detail=f"claim: {block.claim}; {block.detail}"
        ))
    return logs


def write_log(log_entries: list[LogEntry], log_path: Path):
    """Write log entries to JSONL file (append mode)."""
    with log_path.open("a", encoding="utf-8") as f:
        for entry in log_entries:
            f.write(entry.to_json() + "\n")


def clean_markdown(markdown: str, applied_blocks: list[str]) -> str:
    """Strip applied annotation blocks from markdown, wrap in <!-- <applied> --> for history."""
    result = markdown
    for raw_block in applied_blocks:
        # Wrap the entire block in <!-- <applied> -->...<!-- </applied> -->
        wrapped = f"<!-- <applied> -->\n{raw_block}\n<!-- </applied> -->"
        result = result.replace(raw_block, wrapped)
    return result


def main():
    parser = argparse.ArgumentParser(
        description="Apply annotated Markdown revisions to a DOCX in-place"
    )
    parser.add_argument("--md", required=True, help="Path to annotated Markdown file")
    parser.add_argument("--docx", required=True, help="Path to target DOCX file (modified in-place)")
    parser.add_argument("--log", help="Path to JSONL log file (default: .manuscript-revision.log next to DOCX)")
    parser.add_argument("--clean-md", action="store_true",
                        help="Strip applied <updated>/<to-be-removed> blocks from Markdown after success")
    parser.add_argument("--dry-run", action="store_true",
                        help="Parse and validate only, do not modify DOCX")

    args = parser.parse_args()

    md_path = Path(args.md)
    docx_path = Path(args.docx)

    if not md_path.exists():
        print(f"❌ Markdown file not found: {md_path}")
        sys.exit(1)

    if not docx_path.exists():
        print(f"❌ DOCX file not found: {docx_path}")
        sys.exit(1)

    # Default log path
    if args.log:
        log_path = Path(args.log)
    else:
        log_path = docx_path.parent / ".manuscript-revision.log"

    print(f"📄 Reading Markdown: {md_path}")
    print(f"📄 Target DOCX: {docx_path}")
    print(f"📝 Log file: {log_path}")

    markdown = md_path.read_text(encoding="utf-8")

    # Parse all annotation blocks
    updated_blocks, removed_blocks, human_task_blocks, fact_check_blocks = parse_annotation_blocks(markdown)

    print(f"\n🔍 Found annotation blocks:")
    print(f"   <updated>: {len(updated_blocks)}")
    print(f"   <to-be-removed>: {len(removed_blocks)}")
    print(f"   <human-task>: {len(human_task_blocks)}")
    print(f"   <fact-check>: {len(fact_check_blocks)}")

    # Collect all log entries
    all_logs = []

    # Log human tasks and fact checks (these don't modify the docx)
    all_logs.extend(log_human_tasks(human_task_blocks))
    all_logs.extend(log_fact_checks(fact_check_blocks))

    if not args.dry_run:
        # Load DOCX
        print(f"\n📂 Loading DOCX...")
        doc = Document(docx_path)

        # Apply updates
        if updated_blocks:
            print(f"✏️  Applying {len(updated_blocks)} update(s)...")
            all_logs.extend(apply_updated_blocks(doc, updated_blocks, log_path))

        # Apply removals
        if removed_blocks:
            print(f"🗑️  Applying {len(removed_blocks)} removal(s)...")
            all_logs.extend(apply_removed_blocks(doc, removed_blocks, log_path))

        # Save DOCX in-place
        print(f"💾 Saving DOCX in-place...")
        doc.save(docx_path)

        # Optional: clean markdown
        if args.clean_md and (updated_blocks or removed_blocks):
            print(f"🧹 Cleaning Markdown (wrapping applied blocks in <!-- <applied> -->)...")
            applied_raws = [b.raw for b in updated_blocks] + [b.raw for b in removed_blocks]
            cleaned = clean_markdown(markdown, applied_raws)
            md_path.write_text(cleaned, encoding="utf-8")
            print(f"   Updated: {md_path}")

    # Write log
    print(f"📝 Writing log to {log_path}...")
    write_log(all_logs, log_path)

    # Summary
    ok_count = sum(1 for e in all_logs if e.status == "ok")
    failed_count = sum(1 for e in all_logs if e.status == "failed")
    pending_count = sum(1 for e in all_logs if e.status == "pending")

    print(f"\n✅ Done!")
    print(f"   Applied: {ok_count}")
    print(f"   Failed: {failed_count}")
    print(f"   Human tasks pending: {pending_count}")
    print(f"   Fact checks logged: {sum(1 for e in all_logs if e.action == 'fact-check')}")

    if failed_count > 0:
        print(f"\n⚠️  Some operations failed. Check log: {log_path}")
        sys.exit(1)


if __name__ == "__main__":
    main()
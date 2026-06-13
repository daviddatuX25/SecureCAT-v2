#!/usr/bin/env python3
"""Unit tests for md_to_docx.py extended annotation syntax support."""

import tempfile
import json
from pathlib import Path

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))

from md_to_docx import (
    parse_manuscript, 
    apply_annotations, 
    apply_updates,
    find_section_by_heading,
    get_paragraph_in_section,
    replace_paragraph_content,
    delete_specific_paragraph
)
from docx import Document


def test_parse_legacy_tags_only():
    """Test that legacy TAG-based UPDATE/REMOVE parsing still works."""
    md_content = """<!-- META: docx-sha256="abc123" -->

<!-- TAG: C1-01 -->
# Introduction

<!-- UPDATE:START -->
This is updated content.
<!-- UPDATE:END -->

Some original content.

<!-- REMOVE:START -->
This should be removed.
<!-- REMOVE:END -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert meta.get('docx-sha256') == 'abc123'
        assert 'C1-01' in tags
        assert tags['C1-01']['heading'] == 'Introduction'
        assert tags['C1-01']['update'] == 'This is updated content.'
        assert tags['C1-01']['remove'] == 'This should be removed.'
        assert len(annotations) == 0  # No new annotations
        print("✅ test_parse_legacy_tags_only passed")
    finally:
        md_path.unlink()


def test_parse_updated_annotation():
    """Test parsing <updated> annotation blocks."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Some Section

<!-- <updated section="Some Section" para="1">New paragraph content</updated> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'updated'
        assert ann['section'] == 'Some Section'
        assert ann['para'] == 1
        assert ann['content'] == 'New paragraph content'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_updated_annotation passed")
    finally:
        md_path.unlink()


def test_parse_to_be_removed_annotation():
    """Test parsing <to-be-removed> annotation blocks."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Another Section

<!-- <to-be-removed section="Another Section" para="2" reason="Outdated information">Old content to remove</to-be-removed> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'to-be-removed'
        assert ann['section'] == 'Another Section'
        assert ann['para'] == 2
        assert ann['reason'] == 'Outdated information'
        assert ann['content'] == 'Old content to remove'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_to_be_removed_annotation passed")
    finally:
        md_path.unlink()


def test_parse_human_task_annotation():
    """Test parsing <human-task> annotation blocks."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Section

<!-- <human-task type="review" status="pending">Verify citation accuracy</human-task> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'human-task'
        assert ann['task_type'] == 'review'
        assert ann['status'] == 'pending'
        assert ann['content'] == 'Verify citation accuracy'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_human_task_annotation passed")
    finally:
        md_path.unlink()


def test_parse_fact_check_annotation():
    """Test parsing <fact-check> annotation blocks."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Section

<!-- <fact-check claim="User count is 10,000" status="verified">Confirmed via analytics dashboard</fact-check> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'fact-check'
        assert ann['claim'] == 'User count is 10,000'
        assert ann['status'] == 'verified'
        assert ann['content'] == 'Confirmed via analytics dashboard'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_fact_check_annotation passed")
    finally:
        md_path.unlink()


def test_parse_mixed_legacy_and_annotations():
    """Test that both legacy TAG blocks and new annotations can coexist."""
    md_content = """<!-- META: docx-sha256="abc123" -->

<!-- TAG: C1-01 -->
# Introduction

<!-- UPDATE:START -->
Legacy update content.
<!-- UPDATE:END -->

Original intro text.

<!-- <updated section="Introduction" para="1">Annotation update content</updated> -->
<!-- <human-task type="review" status="pending">Check intro flow</human-task> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        # Legacy tag parsed
        assert 'C1-01' in tags
        assert tags['C1-01']['update'] == 'Legacy update content.'
        
        # New annotations parsed
        assert len(annotations) == 2
        assert annotations[0]['type'] == 'updated'
        assert annotations[1]['type'] == 'human-task'
        print("✅ test_parse_mixed_legacy_and_annotations passed")
    finally:
        md_path.unlink()


def test_change_log_aggregation():
    """Test that change log aggregates entries from both legacy tags and annotations."""
    md_content = """<!-- META: docx-sha256="abc123" -->

<!-- TAG: C1-01 -->
# Introduction

<!-- UPDATE:START -->
Legacy update content.
<!-- UPDATE:END -->

<!-- <updated section="Introduction" para="1">Annotation update content</updated> -->
<!-- <human-task type="review" status="pending">Check intro flow</human-task> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        # Create a simple test document
        doc = Document()
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('Original paragraph 1')
        doc.add_paragraph('Original paragraph 2')
        
        change_log = []
        
        # Apply legacy updates
        apply_updates(doc, tags, ['C1-01'], change_log)
        
        # Apply annotations
        apply_annotations(doc, annotations, change_log)
        
        # Verify change log has entries from both sources
        legacy_entries = [e for e in change_log if e['source'] == 'legacy_tag']
        annotation_entries = [e for e in change_log if e['source'] == 'annotation']
        
        assert len(legacy_entries) == 1
        assert legacy_entries[0]['action'] == 'update'
        assert legacy_entries[0]['tag'] == 'C1-01'
        
        assert len(annotation_entries) == 2
        # One update action, one log_only action
        update_entries = [e for e in annotation_entries if e['action'] == 'update']
        log_only_entries = [e for e in annotation_entries if e['action'] == 'log_only']
        assert len(update_entries) == 1
        assert len(log_only_entries) == 1
        assert log_only_entries[0]['type'] == 'human-task'
        
        print("✅ test_change_log_aggregation passed")
    finally:
        md_path.unlink()


def test_find_section_by_heading():
    """Test section locating by heading text."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')
    doc.add_paragraph('Para 2')
    doc.add_heading('Background', level=2)  # Subsection of Introduction
    doc.add_paragraph('Para 3')
    doc.add_paragraph('Para 4')
    doc.add_heading('Methodology', level=1)  # Peer section to Introduction
    doc.add_paragraph('Para 5')
    
    # Find "Introduction" section - should extend to next Heading 1 (Methodology)
    start, end = find_section_by_heading(doc, 'Introduction')
    assert start == 0  # Heading is at index 0
    assert end == 6    # Next Heading 1 (Methodology) at index 6
    
    # Find "Background" section - should be within Introduction, ends at Methodology
    start, end = find_section_by_heading(doc, 'Background')
    assert start == 3  # Heading at index 3
    assert end == 6    # Ends at next Heading 1 (Methodology)
    
    # Find "Methodology" section - extends to end of doc
    start, end = find_section_by_heading(doc, 'Methodology')
    assert start == 6  # Heading at index 6
    assert end == 8    # End of document (8 paragraphs total)
    
    print("✅ test_find_section_by_heading passed")


def test_paragraph_indexing():
    """Test 1-based paragraph indexing within a section."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')  # index 1 within section
    doc.add_paragraph('Para 2')  # index 2 within section
    doc.add_paragraph('Para 3')  # index 3 within section
    
    start, end = find_section_by_heading(doc, 'Introduction')
    
    # para_index=1 should get first paragraph after heading
    p = get_paragraph_in_section(doc, start, end, 1)
    assert p is not None
    assert p.text == 'Para 1'
    
    p = get_paragraph_in_section(doc, start, end, 2)
    assert p is not None
    assert p.text == 'Para 2'
    
    p = get_paragraph_in_section(doc, start, end, 3)
    assert p is not None
    assert p.text == 'Para 3'
    
    p = get_paragraph_in_section(doc, start, end, 4)
    assert p is None  # Out of bounds
    
    # Test replace
    replace_paragraph_content(doc, start, end, 2, 'Replaced content')
    p2 = doc.paragraphs[2]  # heading + para1 + para2 (replaced)
    assert p2.text == 'Replaced content'
    
    print("✅ test_paragraph_indexing passed")


def test_delete_specific_paragraph():
    """Test deleting a specific paragraph by 1-based index."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')
    doc.add_paragraph('Para 2')
    doc.add_paragraph('Para 3')
    
    start, end = find_section_by_heading(doc, 'Introduction')
    
    # Delete para 2
    success = delete_specific_paragraph(doc, start, end, 2)
    assert success
    
    # Verify para 2 is gone
    assert len(doc.paragraphs) == 3  # heading + para1 + para3
    assert doc.paragraphs[1].text == 'Para 1'
    assert doc.paragraphs[2].text == 'Para 3'
    
    print("✅ test_delete_specific_paragraph passed")


def test_apply_updated_annotation():
    """Test applying <updated> annotation to document."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Original paragraph 1')
    doc.add_paragraph('Original paragraph 2')
    
    annotations = [{
        'type': 'updated',
        'section': 'Introduction',
        'para': 1,
        'content': 'Updated via annotation',
        'source': 'annotation'
    }]
    
    change_log = []
    applied = apply_annotations(doc, annotations, change_log)
    
    # Verify paragraph was updated
    assert doc.paragraphs[1].text == 'Updated via annotation'
    assert doc.paragraphs[2].text == 'Original paragraph 2'
    
    # Verify change log
    assert len(change_log) == 1
    assert change_log[0]['source'] == 'annotation'
    assert change_log[0]['type'] == 'updated'
    assert change_log[0]['action'] == 'update'
    assert change_log[0]['section'] == 'Introduction'
    assert change_log[0]['para'] == 1
    
    # Verify applied tracking
    assert 'Introduction' in applied
    assert applied['Introduction'][0]['type'] == 'updated'
    assert applied['Introduction'][0]['para'] == 1
    
    print("✅ test_apply_updated_annotation passed")


def test_apply_to_be_removed_annotation():
    """Test applying <to-be-removed> annotation to document."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')
    doc.add_paragraph('Para 2 - to remove')
    doc.add_paragraph('Para 3')
    
    annotations = [{
        'type': 'to-be-removed',
        'section': 'Introduction',
        'para': 2,
        'reason': 'Outdated',
        'content': 'Para 2 - to remove',
        'source': 'annotation'
    }]
    
    change_log = []
    applied = apply_annotations(doc, annotations, change_log)
    
    # Verify para 2 was removed
    assert len(doc.paragraphs) == 3  # heading + para1 + para3
    assert doc.paragraphs[1].text == 'Para 1'
    assert doc.paragraphs[2].text == 'Para 3'
    
    # Verify change log with reason
    assert len(change_log) == 1
    assert change_log[0]['source'] == 'annotation'
    assert change_log[0]['type'] == 'to-be-removed'
    assert change_log[0]['action'] == 'remove'
    assert change_log[0]['reason'] == 'Outdated'
    assert change_log[0]['section'] == 'Introduction'
    assert change_log[0]['para'] == 2
    
    print("✅ test_apply_to_be_removed_annotation passed")


def test_apply_human_task_and_fact_check_log_only():
    """Test that human-task and fact-check only log, don't modify DOCX."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')
    original_para_count = len(doc.paragraphs)
    
    annotations = [
        {
            'type': 'human-task',
            'task_type': 'review',
            'status': 'pending',
            'content': 'Check citations',
            'source': 'annotation'
        },
        {
            'type': 'fact-check',
            'claim': 'Claim X',
            'status': 'verified',
            'content': 'Verified via source Y',
            'source': 'annotation'
        }
    ]
    
    change_log = []
    applied = apply_annotations(doc, annotations, change_log)
    
    # Verify DOCX unchanged
    assert len(doc.paragraphs) == original_para_count
    assert doc.paragraphs[1].text == 'Para 1'
    
    # Verify change log has both entries
    assert len(change_log) == 2
    log_entries = [e for e in change_log if e['action'] == 'log_only']
    assert len(log_entries) == 2
    assert log_entries[0]['type'] == 'human-task'
    assert log_entries[0]['task_type'] == 'review'
    assert log_entries[1]['type'] == 'fact-check'
    assert log_entries[1]['claim'] == 'Claim X'
    
    # Verify applied tracking is empty for log-only types
    assert applied == {}
    
    print("✅ test_apply_human_task_and_fact_check_log_only passed")


def test_missing_section_error_handling():
    """Test error handling when target section not found."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')
    
    annotations = [{
        'type': 'updated',
        'section': 'Nonexistent Section',
        'para': 1,
        'content': "Won't apply",
        'source': 'annotation'
    }]
    
    change_log = []
    applied = apply_annotations(doc, annotations, change_log)
    
    # Verify error logged
    assert len(change_log) == 1
    assert change_log[0]['action'] == 'error'
    assert 'Section not found' in change_log[0]['details']
    
    # Verify applied tracking empty
    assert applied == {}
    
    print("✅ test_missing_section_error_handling passed")


def test_out_of_bounds_paragraph_error_handling():
    """Test error handling when target paragraph index out of bounds."""
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('Para 1')  # Only 1 paragraph
    
    annotations = [{
        'type': 'updated',
        'section': 'Introduction',
        'para': 5,  # Out of bounds
        'content': "Won't apply",
        'source': 'annotation'
    }]
    
    change_log = []
    applied = apply_annotations(doc, annotations, change_log)
    
    # Verify error logged
    assert len(change_log) == 1
    assert change_log[0]['action'] == 'error'
    assert 'not found' in change_log[0]['details']
    
    print("✅ test_out_of_bounds_paragraph_error_handling passed")


def test_parse_split_updated_annotation():
    """Test parsing <updated> annotation blocks written as split comments."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Some Section

<!-- <updated section="Some Section" para="1"> -->
New paragraph content
<!-- </updated> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'updated'
        assert ann['section'] == 'Some Section'
        assert ann['para'] == 1
        assert ann['content'] == 'New paragraph content'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_split_updated_annotation passed")
    finally:
        md_path.unlink()


def test_parse_split_to_be_removed_annotation():
    """Test parsing <to-be-removed> annotation blocks written as split comments."""
    md_content = """<!-- META: docx-sha256="abc123" -->

# Another Section

<!-- <to-be-removed section="Another Section" para="2" reason="Outdated information"> -->
Old content to remove
<!-- </to-be-removed> -->
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md_content)
        md_path = Path(f.name)
    
    try:
        meta, tags, annotations = parse_manuscript(md_path)
        
        assert len(annotations) == 1
        ann = annotations[0]
        assert ann['type'] == 'to-be-removed'
        assert ann['section'] == 'Another Section'
        assert ann['para'] == 2
        assert ann['reason'] == 'Outdated information'
        assert ann['content'] == 'Old content to remove'
        assert ann['source'] == 'annotation'
        print("✅ test_parse_split_to_be_removed_annotation passed")
    finally:
        md_path.unlink()


def run_all_tests():
    """Run all tests."""
    test_parse_legacy_tags_only()
    test_parse_updated_annotation()
    test_parse_split_updated_annotation()
    test_parse_to_be_removed_annotation()
    test_parse_split_to_be_removed_annotation()
    test_parse_human_task_annotation()
    test_parse_fact_check_annotation()
    test_parse_mixed_legacy_and_annotations()
    test_change_log_aggregation()
    test_find_section_by_heading()
    test_paragraph_indexing()
    test_delete_specific_paragraph()
    test_apply_updated_annotation()
    test_apply_to_be_removed_annotation()
    test_apply_human_task_and_fact_check_log_only()
    test_missing_section_error_handling()
    test_out_of_bounds_paragraph_error_handling()
    print("\n🎉 All tests passed!")


if __name__ == '__main__':
    run_all_tests()

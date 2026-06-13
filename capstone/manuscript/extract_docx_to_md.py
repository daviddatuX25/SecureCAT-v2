#!/usr/bin/env python3
"""
Extract DOCX content to JSON using bookmark boundaries.
Preserves existing MD structure: TAG/UPDATE blocks, annotations, META.

Usage:
  python extract_docx_to_md.py [--docx PATH] [--output PATH]
"""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import argparse
import re
import json

DEFAULT_DOCX = "capstone/output/SecureCAT_Ch1_Ch2_Manuscript[never delete].docx"
DEFAULT_OUTPUT = "capstone/manuscript/extracted_from_docx.json"


def get_bookmarks(doc):
    """Get all bookmarked paragraphs with their tag names and heading info.
    
    IMPORTANT: Deduplicates bookmarks — only the FIRST occurrence of each
    tag name is kept. This prevents corrupt DOCX files with triple-duplicated
    bookmarks from breaking extraction.
    """
    bookmarks = []
    seen_tags = set()

    for i, p in enumerate(doc.paragraphs):
        bookmarks_start = p._p.xpath('.//w:bookmarkStart')
        for bm in bookmarks_start:
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                tag = name[4:]  # remove 'tag_' prefix

                # Skip duplicates — only keep first occurrence
                if tag in seen_tags:
                    continue
                seen_tags.add(tag)

                bookmarks.append({
                    'index': i,
                    'tag': tag,
                    'style': p.style.name,
                    'text': p.text.strip(),
                })
    return bookmarks


def extract_section_content(doc, start_idx, end_idx):
    """Extract content between two paragraph indices (exclusive of start, inclusive of content)."""
    content_parts = []
    for i in range(start_idx + 1, end_idx):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            if p.style.name.startswith('Heading'):
                try:
                    level = int(p.style.name.split()[-1])
                    text = f"{'#' * level} {text}"
                except (ValueError, IndexError):
                    text = f"# {text}"
            content_parts.append(text)
    return '\n\n'.join(content_parts)


def extract_all_sections(doc):
    """Extract all tagged sections from DOCX."""
    bookmarks = get_bookmarks(doc)
    sections = {}

    # Sort by document position
    bookmarks.sort(key=lambda b: b['index'])

    for i, bm in enumerate(bookmarks):
        start = bm['index']
        # End is next bookmark's index or end of doc
        end = bookmarks[i + 1]['index'] if i + 1 < len(bookmarks) else len(doc.paragraphs)

        content = extract_section_content(doc, start, end)
        level_match = re.search(r'(\d+)', bm['style'])
        sections[bm['tag']] = {
            'heading': bm['text'],
            'level': int(level_match.group(1)) if level_match else 2,
            'content': content
        }

    return sections


def main():
    parser = argparse.ArgumentParser(description="Extract DOCX bookmark sections to JSON")
    parser.add_argument('--docx', default=DEFAULT_DOCX, help='Path to DOCX file')
    parser.add_argument('--output', default=DEFAULT_OUTPUT, help='Path to output JSON')
    args = parser.parse_args()

    docx_path = Path(args.docx)
    output_path = Path(args.output)

    if not docx_path.exists():
        print(f"❌ DOCX not found: {docx_path}")
        return 1

    doc = Document(docx_path)

    # Extract bookmarked sections
    sections = extract_all_sections(doc)

    print(f"Found {len(sections)} bookmarked sections:")
    for tag, data in sections.items():
        preview = data['content'][:100] + '...' if len(data['content']) > 100 else data['content']
        print(f"  {tag}: level={data['level']}, heading='{data['heading']}', "
              f"content_len={len(data['content'])}, preview='{preview}'")

    # Save for inspection
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(sections, f, indent=2, ensure_ascii=False)
    print(f"\nSaved to {output_path}")

    return 0


if __name__ == '__main__':
    exit(main())
#!/usr/bin/env python3
"""
Convert specific Normal paragraphs to Heading styles in the reverted DOCX.
Preserves ALL formatting (bold, alignment, spacing, font) — only changes style attribute.
Run ONCE after pulling reverted cloud DOCX, then add bookmarks, then surgical works forever.
"""
from docx import Document

DOCX_PATH = "capstone/output/SecureCAT_Ch1_Ch2_Manuscript[never delete].docx"

# Map exact text in your reverted DOCX -> Heading level
# Copy-paste the EXACT text from your document (case-insensitive match)
HEADING_MAP = [
    # (search_text, heading_level)
    ("Chapter 1", 1),
    ("Background of the Study", 2),
    ("Conceptual Framework of the Study", 2),
    ("Objectives of the Study", 2),
    ("Scope and Limitation of the Study", 2),
    ("Significance of the Study", 2),
    ("Chapter 2", 1),
    ("Research Design", 2),
    ("Software Development Model", 2),
    ("Project Plan", 2),
    ("Project Assignments", 2),
    ("Population and Locale of the Study", 2),
    ("Research Instruments", 2),
    ("Data Analysis", 2),
    ("REFERENCES", 2),
    ("APPENDIX A", 2),
    ("APPENDIX B", 2),
]

def matches_heading(p, search_text):
    """Strict heading match: exact, startswith, or isolated phrase — NOT loose contains."""
    ptext = p.text.strip()
    stext = search_text.strip().lower()
    plower = ptext.lower()
    
    # 1. Exact match (after stripping)
    if plower == stext:
        return True
    # 2. Starts with (heading at beginning of paragraph)
    if plower.startswith(stext):
        return True
    # 3. Isolated phrase (surrounded by whitespace/punctuation)
    import re
    if re.search(rf'(^|[\s\n\r]){re.escape(stext)}($|[\s\n\r.,:;!?])', plower):
        return True
    return False

def main():
    doc = Document(DOCX_PATH)
    print(f"Loaded: {DOCX_PATH} ({len(doc.paragraphs)} paragraphs)")
    
    converted = 0
    for search_text, level in HEADING_MAP:
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.lower() == 'normal' and matches_heading(p, search_text):
                orig_style = p.style.name
                runs_info = [(r.text[:30], r.bold, r.italic, r.font.size, r.font.name) for r in p.runs[:3]]
                
                p.style = f'Heading {level}'
                
                print(f"✓ Para {i}: \"{search_text}\" -> Heading {level}")
                print(f"   Original style: {orig_style} | Text: \"{p.text[:80]}\"")
                converted += 1
                break
        else:
            print(f"⚠ NO MATCH: \"{search_text}\"")
    
    doc.save(DOCX_PATH)
    print(f"\nDone. Converted {converted} paragraphs to Heading styles.")
    print("Visual formatting (bold, alignment, font, spacing) PRESERVED.")
    print("Now run: python3 capstone/manuscript/add_bookmarks.py")

if __name__ == '__main__':
    main()
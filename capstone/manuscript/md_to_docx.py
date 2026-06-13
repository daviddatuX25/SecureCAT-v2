#!/usr/bin/env python3
"""
Core MD → DOCX converter with tag-based surgical updates.
Skips Mermaid/diagram auto-generation (diagrams are human tasks).
Uses template table formatting (no left/vertical middle borders).

Extended to support HTML-comment annotation syntax for review-gated surgical edits:
- <updated section="X" para="Y">paragraph content</updated>
- <to-be-removed section="X" para="Y" reason="...">paragraph content</to-be-removed>
- <human-task type="..." status="...">task description</human-task>
- <fact-check claim="..." status="...">verification notes</fact-check>

Includes --dry-run / --review mode for previewing changes before applying.
"""
import argparse
import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ======== MANUSCRIPT PARSING ========

def parse_manuscript(md_path: Path) -> tuple[dict, dict, list]:
    """Parse manuscript MD for META, TAGs, UPDATE/REMOVE blocks, and new annotation blocks.

    Returns:
        tuple: (meta, tags, annotations)
        - meta: dict from META tag
        - tags: dict of legacy TAG-based sections with update/remove/raw_content
        - annotations: list of dict for new HTML-comment annotation blocks
    """
    content = md_path.read_text(encoding='utf-8')

    # Parse META tag
    meta = {}
    meta_match = re.search(r'<!--\s*META:\s*(.*?)\s*-->', content)
    if meta_match:
        for part in meta_match.group(1).split():
            if '=' in part:
                k, v = part.split('=', 1)
                meta[k] = v.strip('"')

    # Parse TAGs with UPDATE/REMOVE blocks (legacy syntax)
    tags = {}
    tag_pattern = re.compile(r'<!--\s*TAG:\s*(\S+)\s*-->')
    tag_positions = [(m.start(), m.group(1)) for m in tag_pattern.finditer(content)]

    for i, (pos, tag_name) in enumerate(tag_positions):
        end_pos = tag_positions[i + 1][0] if i + 1 < len(tag_positions) else len(content)
        section_content = content[pos:end_pos]

        update_match = re.search(
            r'<!--\s*UPDATE:START\s*-->(.*?)<!--\s*UPDATE:END\s*-->',
            section_content, re.DOTALL
        )
        update_content = update_match.group(1).strip() if update_match else None

        remove_match = re.search(
            r'<!--\s*REMOVE:START\s*-->(.*?)<!--\s*REMOVE:END\s*-->',
            section_content, re.DOTALL
        )
        remove_content = remove_match.group(1).strip() if remove_match else None

        section_no_tag = re.sub(r'<!--\s*TAG:\s*\S+\s*-->\s*', '', section_content, count=1)
        heading_match = re.search(r'^(#{1,4})\s+(.+)$', section_no_tag, re.MULTILINE)
        heading = heading_match.group(2).strip() if heading_match else tag_name
        heading_level = len(heading_match.group(1)) if heading_match else 2

        raw_content = section_no_tag
        raw_content = re.sub(r'<!--\s*UPDATE:START\s*-->.*?<!--\s*UPDATE:END\s*-->', '', raw_content, flags=re.DOTALL)
        raw_content = re.sub(r'<!--\s*REMOVE:START\s*-->.*?<!--\s*REMOVE:END\s*-->', '', raw_content, flags=re.DOTALL)
        raw_content = raw_content.strip()

        tags[tag_name] = {
            'heading': heading,
            'level': heading_level,
            'update': update_content,
            'remove': remove_content,
            'raw_content': raw_content
        }

    # Parse new HTML-comment annotation blocks
    annotations = []

    # Pattern for <updated section="X" para="Y">content</updated> (both single and split comments)
    updated_pattern = re.compile(
        r'<!--\s*<updated\s+section="([^"]+)"\s+para="(\d+)"\s*>\s*-->(.*?)<!--\s*</updated>\s*-->'
        r'|'
        r'<!--\s*<updated\s+section="([^"]+)"\s+para="(\d+)"\s*>(.*?)</updated>\s*-->',
        re.DOTALL
    )
    for m in updated_pattern.finditer(content):
        if m.group(1) is not None:
            section = m.group(1)
            para = int(m.group(2))
            val = m.group(3)
        else:
            section = m.group(4)
            para = int(m.group(5))
            val = m.group(6)
        annotations.append({
            'type': 'updated',
            'section': section,
            'para': para,
            'content': val.strip(),
            'source': 'annotation'
        })

    # Pattern for <to-be-removed section="X" para="Y" reason="...">content</to-be-removed> (both single and split comments)
    removed_pattern = re.compile(
        r'<!--\s*<to-be-removed\s+section="([^"]+)"\s+para="(\d+)"\s+reason="([^"]*)"\s*>\s*-->(.*?)<!--\s*</to-be-removed>\s*-->'
        r'|'
        r'<!--\s*<to-be-removed\s+section="([^"]+)"\s+para="(\d+)"\s+reason="([^"]*)"\s*>(.*?)</to-be-removed>\s*-->',
        re.DOTALL
    )
    for m in removed_pattern.finditer(content):
        if m.group(1) is not None:
            section = m.group(1)
            para = int(m.group(2))
            reason = m.group(3)
            val = m.group(4)
        else:
            section = m.group(5)
            para = int(m.group(6))
            reason = m.group(7)
            val = m.group(8)
        annotations.append({
            'type': 'to-be-removed',
            'section': section,
            'para': para,
            'reason': reason.strip(),
            'content': val.strip(),
            'source': 'annotation'
        })

    # Pattern for <human-task type="..." status="...">content</human-task> (both single and split comments)
    human_task_pattern = re.compile(
        r'<!--\s*<human-task\s+type="([^"]+)"\s+status="([^"]+)"\s*>\s*-->(.*?)<!--\s*</human-task>\s*-->'
        r'|'
        r'<!--\s*<human-task\s+type="([^"]+)"\s+status="([^"]+)"\s*>(.*?)</human-task>\s*-->',
        re.DOTALL
    )
    for m in human_task_pattern.finditer(content):
        if m.group(1) is not None:
            task_type = m.group(1)
            status = m.group(2)
            val = m.group(3)
        else:
            task_type = m.group(4)
            status = m.group(5)
            val = m.group(6)
        annotations.append({
            'type': 'human-task',
            'task_type': task_type,
            'status': status,
            'content': val.strip(),
            'source': 'annotation'
        })

    # Pattern for <fact-check claim="..." status="...">content</fact-check> (both single and split comments)
    fact_check_pattern = re.compile(
        r'<!--\s*<fact-check\s+claim="([^"]+)"\s+status="([^"]+)"\s*>\s*-->(.*?)<!--\s*</fact-check>\s*-->'
        r'|'
        r'<!--\s*<fact-check\s+claim="([^"]+)"\s+status="([^"]+)"\s*>(.*?)</fact-check>\s*-->',
        re.DOTALL
    )
    for m in fact_check_pattern.finditer(content):
        if m.group(1) is not None:
            claim = m.group(1)
            status = m.group(2)
            val = m.group(3)
        else:
            claim = m.group(4)
            status = m.group(5)
            val = m.group(6)
        annotations.append({
            'type': 'fact-check',
            'claim': claim,
            'status': status,
            'content': val.strip(),
            'source': 'annotation'
        })

    return meta, tags, annotations


# ======== DOCX HELPERS ========

def sanitize_bookmark(name: str) -> str:
    return re.sub(r'[^a-zA-Z0-9_]', '_', name)


def get_or_create_bookmark(doc: Document, tag_name: str, heading_text: str, heading_level: int) -> str:
    bookmark_name = f"tag_{sanitize_bookmark(tag_name)}"

    # 1. Check if the bookmark already exists in the document.
    for i, p in enumerate(doc.paragraphs):
        if p._p.xpath(f'.//w:bookmarkStart[@w:name="{bookmark_name}"]'):
            return bookmark_name

    # 2. If it doesn't exist, locate the best paragraph to place it.
    heading_text_clean = heading_text.strip().lower()
    
    # Try exact match on Heading style paragraphs first
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().lower() == heading_text_clean:
            create_bookmark_at_paragraph(p, bookmark_name)
            return bookmark_name

    # Try substring match on Heading style paragraphs
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and heading_text_clean in p.text.strip().lower():
            create_bookmark_at_paragraph(p, bookmark_name)
            return bookmark_name

    # Try exact match on any paragraph (styled Normal or otherwise)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == heading_text_clean:
            create_bookmark_at_paragraph(p, bookmark_name)
            return bookmark_name

    # Try fuzzy exact matches (e.g. plural vs singular, "and limitation" vs "and delimitations")
    def normalize_text(t):
        t = re.sub(r'[^a-z0-9\s]', '', t.lower())
        t = t.replace('delimitations', 'limitation').replace('delimitations', 'limitations')
        t = t.replace('limitations', 'limitation')
        t = t.replace('assignments', 'assignment')
        t = t.replace('significance', 'importance')
        t = t.replace('importance', 'significance')
        return ' '.join(t.split())

    norm_heading = normalize_text(heading_text_clean)
    for i, p in enumerate(doc.paragraphs):
        norm_p = normalize_text(p.text)
        if norm_p == norm_heading or (norm_heading in norm_p and len(norm_p) < 150):
            create_bookmark_at_paragraph(p, bookmark_name)
            return bookmark_name

    # Fallback: create at end
    p = doc.paragraphs[-1]
    existing = p._p.xpath(f'.//w:bookmarkStart[@w:name="{bookmark_name}"]')
    if not existing:
        create_bookmark_at_paragraph(p, bookmark_name)
    return bookmark_name


def create_bookmark_at_paragraph(paragraph, bookmark_name: str):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(hash(bookmark_name) % 1000000))
    start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(hash(bookmark_name) % 1000000))
    end.set(qn('w:name'), bookmark_name)
    paragraph._p.append(end)


def find_section_paragraphs(doc: Document, bookmark_name: str, tag_data: dict) -> tuple[int, int]:
    """Find the start and end paragraph indices for a section identified by bookmark.

    Section boundaries are determined by the NEXT tag_ bookmark in document order.
    This avoids the old bug where stopping at 'any bookmarked heading' caused sections
    to end immediately when all headings are bookmarked.
    """
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p._p.xpath(f'.//w:bookmarkStart[@w:name="{bookmark_name}"]'):
            start_idx = i
            break

    if start_idx == -1:
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading') and tag_data['heading'].lower() in p.text.lower():
                start_idx = i
                break

    if start_idx == -1:
        return (-1, -1)

    # Find ALL tag_ bookmarks in document order (deduplicated)
    seen_tags = set()
    all_bookmarks = []
    for i, p in enumerate(doc.paragraphs):
        for bm in p._p.xpath('.//w:bookmarkStart'):
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_') and name not in seen_tags:
                seen_tags.add(name)
                all_bookmarks.append(i)
    all_bookmarks.sort()

    # End index is the next bookmark AFTER start_idx
    end_idx = len(doc.paragraphs)
    for bm_idx in all_bookmarks:
        if bm_idx > start_idx:
            end_idx = bm_idx
            break

    return (start_idx, end_idx)


def find_section_by_heading(doc: Document, heading_text: str) -> tuple[int, int]:
    """Find section paragraph indices by heading text (for annotation-based operations)."""
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and heading_text.lower() in p.text.lower():
            start_idx = i
            break

    if start_idx == -1:
        return (-1, -1)

    current_level = 0
    for p in doc.paragraphs[start_idx:start_idx+1]:
        if p.style.name.startswith('Heading'):
            try:
                current_level = int(p.style.name[-1])
            except:
                current_level = 2

    end_idx = len(doc.paragraphs)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading'):
            try:
                level = int(p.style.name[-1])
            except:
                level = 2
            if level <= current_level:
                end_idx = i
                break

    return (start_idx, end_idx)


def get_paragraph_in_section(doc: Document, start_idx: int, end_idx: int, para_index: int):
    """Get the paragraph at 1-based index within a section (excluding the heading)."""
    if start_idx < 0 or end_idx <= start_idx:
        return None

    target_idx = start_idx + para_index
    if target_idx >= end_idx or target_idx >= len(doc.paragraphs):
        return None

    return doc.paragraphs[target_idx]


def delete_specific_paragraph(doc: Document, start_idx: int, end_idx: int, para_index: int):
    """Delete a specific paragraph by 1-based index within a section."""
    p = get_paragraph_in_section(doc, start_idx, end_idx, para_index)
    if p is not None:
        p._element.getparent().remove(p._element)
        return True
    return False


def replace_paragraph_content(doc: Document, start_idx: int, end_idx: int, para_index: int, new_content: str):
    """Replace content of a specific paragraph by 1-based index within a section."""
    p = get_paragraph_in_section(doc, start_idx, end_idx, para_index)
    if p is not None:
        # Clear existing runs and add new content
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new_content
        else:
            p.add_run(new_content)
        return True
    return False


def delete_paragraphs(doc: Document, start_idx: int, end_idx: int, remove_text: str = None):
    if start_idx < 0 or end_idx <= start_idx:
        return

    to_delete = []
    for i in range(start_idx + 1, end_idx):
        p = doc.paragraphs[i]
        
        # Check if the paragraph contains a page break or section properties
        has_page_break = False
        if p._element.xpath('.//w:pageBreakBefore') or p._element.xpath('.//w:br[@w:type="page"]') or p._element.xpath('.//w:sectPr'):
            has_page_break = True

        if remove_text:
            if remove_text.strip() in p.text:
                if has_page_break:
                    # Clear text in runs that don't contain the break
                    for run in p.runs:
                        if not run._element.xpath('.//w:br[@w:type="page"]'):
                            run.text = ""
                else:
                    to_delete.append(i)
        else:
            if has_page_break:
                # Clear text in runs that don't contain the break
                for run in p.runs:
                    if not run._element.xpath('.//w:br[@w:type="page"]'):
                        run.text = ""
            else:
                to_delete.append(i)

    for i in reversed(to_delete):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def insert_section_content(doc: Document, start_idx: int, end_idx: int, new_content: str, heading_level: int, is_references: bool = False):
    if start_idx < 0:
        start_idx = len(doc.paragraphs) - 1

    delete_paragraphs(doc, start_idx, end_idx)

    # Get the heading paragraph's XML element for direct insertion
    heading_element = doc.paragraphs[start_idx]._p

    lines = new_content.split('\n')
    # Process lines in REVERSE so each insert goes right after heading
    # (inserting A, B, C after heading in reverse gives: heading → A → B → C)
    non_empty_lines = [line.strip() for line in lines if line.strip()]

    # Insert in reverse order using addnext (each goes right after heading)
    for line in reversed(non_empty_lines):
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('# ').strip()
            p = doc.add_paragraph(text, style=f'Heading {min(level, 4)}')
            for run in p.runs:
                run.font.name = 'Times New Roman'
        else:
            p = doc.add_paragraph(line, style='Normal')
            if is_references:
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.left_indent = Inches(0)
            
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if not p.runs:
                p.add_run()
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        # Move from end-of-document to right after heading
        heading_element.addnext(p._p)


def move_paragraph_after(paragraph, after_paragraph):
    after_paragraph._p.addnext(paragraph._p)


def extract_existing_text(doc: Document, start_idx: int, end_idx: int) -> str:
    parts = []
    for i in range(start_idx + 1, end_idx):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def normalize_comparable_text(text: str) -> str:
    if not text:
        return ""
    # Remove HTML comments, markdown headings, and normalize spaces/newlines
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()


def apply_updates(doc: Document, tags: dict, update_tags: list = None, change_log: list = None) -> dict:
    """Apply legacy TAG-based updates to the document."""
    if update_tags is None:
        update_tags = list(tags.keys())

    # Pass 1: Ensure all bookmarks are created first at their initial positions.
    for tag_name in update_tags:
        if tag_name not in tags:
            continue
        tag_data = tags[tag_name]
        get_or_create_bookmark(doc, tag_name, tag_data['heading'], tag_data['level'])

    tag_map = {}

    # Pass 2: Apply updates (delete and insert) using the now stable bookmarks.
    for tag_name in update_tags:
        if tag_name not in tags:
            print(f"⚠️  Tag '{tag_name}' not found in manuscript")
            continue

        tag_data = tags[tag_name]
        bookmark_name = f"tag_{sanitize_bookmark(tag_name)}"
        start_idx, end_idx = find_section_paragraphs(doc, bookmark_name, tag_data)

        tag_map[tag_name] = {
            'bookmark': bookmark_name,
            'start_idx': start_idx,
            'end_idx': end_idx,
            'heading': tag_data['heading'],
            'level': tag_data['level']
        }

        if tag_data['remove'] and start_idx >= 0:
            delete_paragraphs(doc, start_idx, end_idx, tag_data['remove'])
            print(f"  REMOVED content for {tag_name}")
            if change_log is not None:
                change_log.append({
                    'source': 'legacy_tag',
                    'tag': tag_name,
                    'action': 'remove',
                    'section': tag_data['heading'],
                    'details': tag_data['remove']
                })

        if tag_data['update'] and start_idx >= 0:
            existing_text = extract_existing_text(doc, start_idx, end_idx)
            if normalize_comparable_text(tag_data['update']) == normalize_comparable_text(existing_text):
                print(f"  SKIPPED update for {tag_name} (content matches)")
            else:
                insert_section_content(doc, start_idx, end_idx, tag_data['update'], tag_data['level'], is_references=(tag_name == 'references_list'))
                print(f"  UPDATED content for {tag_name}")
                if change_log is not None:
                    change_log.append({
                        'source': 'legacy_tag',
                        'tag': tag_name,
                        'action': 'update',
                        'section': tag_data['heading'],
                        'details': tag_data['update'][:100]
                    })

        if not tag_data['update'] and not tag_data['remove']:
            print(f"  No changes for {tag_name}")

    return tag_map


def apply_annotations(doc: Document, annotations: list, change_log: list = None) -> dict:
    """Apply new HTML-comment annotation-based surgical edits to the document.

    Returns a dict mapping section headings to applied annotation info.
    """
    applied = {}

    for ann in annotations:
        ann_type = ann['type']

        if ann_type in ('updated', 'to-be-removed'):
            # These modify the DOCX - find section by heading, then target paragraph
            section_heading = ann['section']
            para_index = ann['para']

            start_idx, end_idx = find_section_by_heading(doc, section_heading)

            if start_idx == -1:
                print(f"⚠️  Section '{section_heading}' not found for annotation {ann_type}")
                if change_log is not None:
                    change_log.append({
                        'source': 'annotation',
                        'type': ann_type,
                        'section': section_heading,
                        'para': para_index,
                        'action': 'error',
                        'details': f'Section not found: {section_heading}'
                    })
                continue

            if ann_type == 'updated':
                # Replace paragraph content
                success = replace_paragraph_content(doc, start_idx, end_idx, para_index, ann['content'])
                if success:
                    print(f"  UPDATED paragraph {para_index} in section '{section_heading}' via annotation")
                    if change_log is not None:
                        change_log.append({
                            'source': 'annotation',
                            'type': 'updated',
                            'section': section_heading,
                            'para': para_index,
                            'action': 'update',
                            'details': ann['content'][:100]
                        })
                    applied.setdefault(section_heading, []).append({
                        'type': 'updated',
                        'para': para_index,
                        'content': ann['content']
                    })
                else:
                    print(f"⚠️  Paragraph {para_index} not found in section '{section_heading}' for update")
                    if change_log is not None:
                        change_log.append({
                            'source': 'annotation',
                            'type': 'updated',
                            'section': section_heading,
                            'para': para_index,
                            'action': 'error',
                            'details': f'Paragraph {para_index} not found'
                        })

            elif ann_type == 'to-be-removed':
                # Remove paragraph and log reason
                success = delete_specific_paragraph(doc, start_idx, end_idx, para_index)
                if success:
                    print(f"  REMOVED paragraph {para_index} in section '{section_heading}' via annotation (reason: {ann['reason']})")
                    if change_log is not None:
                        change_log.append({
                            'source': 'annotation',
                            'type': 'to-be-removed',
                            'section': section_heading,
                            'para': para_index,
                            'action': 'remove',
                            'reason': ann['reason'],
                            'details': ann['content'][:100]
                        })
                    applied.setdefault(section_heading, []).append({
                        'type': 'to-be-removed',
                        'para': para_index,
                        'reason': ann['reason'],
                        'content': ann['content']
                    })
                else:
                    print(f"⚠️  Paragraph {para_index} not found in section '{section_heading}' for removal")
                    if change_log is not None:
                        change_log.append({
                            'source': 'annotation',
                            'type': 'to-be-removed',
                            'section': section_heading,
                            'para': para_index,
                            'action': 'error',
                            'details': f'Paragraph {para_index} not found; reason: {ann["reason"]}'
                        })

        elif ann_type in ('human-task', 'fact-check'):
            # These do NOT modify DOCX - only add to change log
            if change_log is not None:
                entry = {
                    'source': 'annotation',
                    'type': ann_type,
                    'action': 'log_only',
                }
                if ann_type == 'human-task':
                    entry.update({
                        'task_type': ann['task_type'],
                        'status': ann['status'],
                        'details': ann['content']
                    })
                else:  # fact-check
                    entry.update({
                        'claim': ann['claim'],
                        'status': ann['status'],
                        'details': ann['content']
                    })
                change_log.append(entry)
                print(f"  LOGGED {ann_type}: {entry.get('task_type') or entry.get('claim')}")

    return applied


# ======== TABLE FORMATTING (Template Style) ========

def format_table_template_style(table):
    """Apply template table style: no left/vertical middle borders, centered alignment."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders first, then add only needed ones
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Remove existing borders
    for borders in tblPr.xpath('.//w:tblBorders'):
        tblPr.remove(borders)

    # Add borders: only top, bottom, and horizontal inside (no left/right/vertical middle)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'bottom', 'insideH']:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')  # 0.5pt
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    # insideV, left, right are intentionally omitted
    tblPr.append(borders)

    # Header row bold
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'

    # First row (header) bold
    if table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def create_table_from_markdown(doc: Document, markdown_table: str):
    """Create a table from markdown pipe syntax with template formatting."""
    lines = [l.strip() for l in markdown_table.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    # Parse header
    header_cells = [c.strip() for c in lines[0].split('|') if c.strip()]
    num_cols = len(header_cells)

    # Parse rows
    rows_data = []
    for line in lines[2:]:  # Skip separator line
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if len(cells) == num_cols:
            rows_data.append(cells)

    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    table.style = 'Table Grid'

    # Header
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = cell_text

    # Data rows
    for r, row_data in enumerate(rows_data):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = cell_text

    format_table_template_style(table)
    return table


# ======== HIGH-LEVEL OPERATIONS ========

def full_rebuild(template_path: Path, tags: dict, output_path: Path):
    if template_path.exists():
        doc = Document(template_path)
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        doc = Document()

    for tag_name, tag_data in tags.items():
        p = doc.add_paragraph(tag_data['heading'], style=f'Heading {tag_data["level"]}')
        create_bookmark_at_paragraph(p, f"tag_{sanitize_bookmark(tag_name)}")

        content = tag_data['update'] or tag_data['raw_content']
        if content:
            insert_section_content(doc, len(doc.paragraphs) - 1, len(doc.paragraphs), content, tag_data['level'])

    doc.save(output_path)
    print(f"✅ Full rebuild saved to {output_path}")


def compute_docx_sha256(docx_path: Path) -> str:
    return hashlib.sha256(docx_path.read_bytes()).hexdigest()


def update_manuscript_meta(md_path: Path, new_meta: dict):
    content = md_path.read_text(encoding='utf-8')
    meta_str = ' '.join(f'{k}="{v}"' for k, v in new_meta.items())
    new_content = re.sub(
        r'<!--\s*META:.*?-->',
        f'<!-- META: {meta_str} -->',
        content
    )
    if '<!-- META:' not in content:
        new_content = f'<!-- META: {meta_str} -->\n\n{new_content}'
    md_path.write_text(new_content, encoding='utf-8')


def save_tag_map(tag_map: dict, json_path: Path):
    data = {
        'version': 1,
        'generated': datetime.now(timezone.utc).isoformat(),
        'tags': {k: {kk: vv for kk, vv in v.items() if kk != 'paragraph_indices'} for k, v in tag_map.items()}
    }
    json_path.write_text(json.dumps(data, indent=2))


def load_tag_map(json_path: Path) -> dict:
    if json_path.exists():
        return json.loads(json_path.read_text())
    return {}


# ======== CLI ========

def main():
    parser = argparse.ArgumentParser(
        description="MD → DOCX converter with tag-based surgical updates and HTML-comment annotation support. "
                    "Skips Mermaid/diagrams (human tasks). Template table formatting."
    )
    parser.add_argument('--md', required=True, help="Manuscript MD path")
    parser.add_argument('--template', help="Template DOCX path (for full rebuild)")
    parser.add_argument('--docx', help="Master DOCX path (for surgical update)")
    parser.add_argument('--output', required=True, help="Output DOCX path")
    parser.add_argument('--full-rebuild', action='store_true', help="Full rebuild from template")
    parser.add_argument('--update-tags', nargs='+', help="Specific tags to update")
    parser.add_argument('--check-drift', action='store_true', help="Check drift only")
    parser.add_argument('--dry-run', '--review', action='store_true',
                        help="Preview all changes without modifying DOCX or writing output. "
                             "Shows what would be updated/removed/logged.")
    args = parser.parse_args()

    md_path = Path(args.md)
    output_path = Path(args.output)

    meta, tags, annotations = parse_manuscript(md_path)
    print(f"Parsed {len(tags)} tags and {len(annotations)} annotations from {md_path}")

    # DRY RUN CHECK - must be FIRST, before any drift checks or processing
    if args.dry_run:
        print("\n" + "="*60)
        print("DRY RUN / REVIEW MODE — No changes will be applied")
        print("="*60)

        # Show legacy tag-based changes
        if args.full_rebuild:
            print("\n📋 FULL REBUILD (from template)")
            for tag_name, tag_data in tags.items():
                if tag_data['update']:
                    print(f"  🔄 UPDATE tag '{tag_name}' (section: {tag_data['heading']})")
                    print(f"      Content preview: {tag_data['update'][:150]}...")
                if tag_data['remove']:
                    print(f"  🗑️  REMOVE tag '{tag_name}' (section: {tag_data['heading']})")
                    print(f"      Content preview: {tag_data['remove'][:150]}...")
        else:
            # Surgical update mode
            print("\n📋 SURGICAL UPDATE (legacy tags)")
            update_tags = args.update_tags if args.update_tags else list(tags.keys())
            for tag_name in update_tags:
                if tag_name in tags:
                    tag_data = tags[tag_name]
                    if tag_data['update']:
                        print(f"  🔄 UPDATE tag '{tag_name}' (section: {tag_data['heading']})")
                        print(f"      Content preview: {tag_data['update'][:150]}...")
                    if tag_data['remove']:
                        print(f"  🗑️  REMOVE tag '{tag_name}' (section: {tag_data['heading']})")
                        print(f"      Content preview: {tag_data['remove'][:150]}...")

        # Show annotation-based surgical edits
        surgical_annotations = [a for a in annotations if a['type'] in ('updated', 'to-be-removed')]
        log_only_annotations = [a for a in annotations if a['type'] in ('human-task', 'fact-check')]

        print("\n📋 ANNOTATION-BASED SURGICAL EDITS")
        if surgical_annotations:
            for ann in surgical_annotations:
                if ann['type'] == 'updated':
                    print(f"  🔄 UPDATE paragraph {ann['para']} in section '{ann['section']}'")
                    print(f"      New content: {ann['content'][:150]}...")
                elif ann['type'] == 'to-be-removed':
                    print(f"  🗑️  REMOVE paragraph {ann['para']} in section '{ann['section']}'")
                    print(f"      Reason: {ann['reason']}")
                    print(f"      Current content: {ann['content'][:150]}...")
        else:
            print("  (none)")

        print("\n📋 ANNOTATIONS (LOG ONLY — no DOCX modification)")
        if log_only_annotations:
            for ann in log_only_annotations:
                if ann['type'] == 'human-task':
                    print(f"  📝 HUMAN TASK [{ann['task_type']}]: {ann['content'][:100]}...")
                elif ann['type'] == 'fact-check':
                    print(f"  ⚠️  FACT CHECK [{ann['status']}]: Claim='{ann['claim'][:80]}...' Notes: {ann['content'][:80]}...")
        else:
            print("  (none)")

        # Drift check preview
        if not args.full_rebuild and args.docx:
            docx_path = Path(args.docx)
            if docx_path.exists():
                current_sha = compute_docx_sha256(docx_path)
                expected_sha = meta.get('docx-sha256')
                if expected_sha:
                    status = "✅ MATCH" if current_sha == expected_sha else "⚠️  MISMATCH (drift)"
                    print(f"\n📋 DRIFT CHECK: {status}")
                    print(f"  Expected: {expected_sha}")
                    print(f"  Actual:   {current_sha}")
                else:
                    print(f"\n📋 DRIFT CHECK: No META SHA256 found — would need full rebuild first")

        print("\n" + "="*60)
        print("DRY RUN COMPLETE — No files modified, no output written")
        print("="*60)
        return 0

    if args.check_drift:
        if not args.docx:
            print("❌ --docx required for drift check")
            return 2
        docx_path = Path(args.docx)
        if not docx_path.exists():
            print("❌ DOCX not found")
            return 2
        current_sha = compute_docx_sha256(docx_path)
        expected_sha = meta.get('docx-sha256')
        if expected_sha and current_sha == expected_sha:
            print(f"✅ PASS: DOCX matches META")
            return 0
        else:
            print(f"❌ DRIFT: expected {expected_sha}, got {current_sha}")
            return 1

    change_log = []

    if args.full_rebuild:
        if not args.template:
            print("❌ --template required for full rebuild")
            return 2
        full_rebuild(Path(args.template), tags, output_path)
    else:
        if not args.docx:
            print("❌ --docx required for surgical update")
            return 2
        docx_path = Path(args.docx)

        if 'docx-sha256' in meta:
            current_sha = compute_docx_sha256(docx_path) if docx_path.exists() else ""
            if current_sha and current_sha != meta['docx-sha256']:
                print(f"⚠️  DRIFT DETECTED: DOCX SHA256 mismatch!")
                print(f"   Expected: {meta['docx-sha256']}")
                print(f"   Actual:   {current_sha}")
                return 1

        from docx import Document
        if docx_path.exists():
            doc = Document(docx_path)
        elif args.template:
            doc = Document(Path(args.template))
        else:
            doc = Document()

        # Apply legacy tag-based updates
        tag_map = apply_updates(doc, tags, args.update_tags, change_log)

        # Apply new annotation-based surgical edits
        applied_annotations = apply_annotations(doc, annotations, change_log)

        doc.save(output_path)
        save_tag_map(tag_map, output_path.parent / 'tag_map.json')

        # Save change log
        change_log_path = output_path.parent / 'change_log.json'
        change_log_data = {
            'version': 1,
            'generated': datetime.now(timezone.utc).isoformat(),
            'entries': change_log
        }
        change_log_path.write_text(json.dumps(change_log_data, indent=2))
        print(f"📝 Change log saved to {change_log_path}")

    new_sha = compute_docx_sha256(output_path)
    new_meta = {
        'docx-sync-version': datetime.now(timezone.utc).isoformat(),
        'docx-sha256': new_sha
    }
    update_manuscript_meta(md_path, new_meta)
    print(f"✅ Done. DOCX: {output_path}, SHA256: {new_sha}")
    return 0


if __name__ == '__main__':
    sys.exit(main())
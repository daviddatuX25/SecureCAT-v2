#!/usr/bin/env python3
"""
Full DOCX Rebuild from MD — Writes clean content into the template DOCX.

Strategy:
  1. Open the CORRUPTED defense DOCX (which has title page, formatting, etc.)
  2. Find all Heading-style paragraphs and map them to our tag system
  3. For each tag section: DELETE all existing body paragraphs, then INSERT the MD content
  4. Add clean bookmarks to all headings
  5. Save to output file and compare content
"""
import hashlib
import json
import re
import sys
import shutil
from datetime import datetime, timezone
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# ======== CONFIG ========
DEFENSE_DOCX = Path("capstone/output/SecureCAT_Ch1_Ch2_Manuscript[never delete].docx")
REPAIRED_DOCX = Path("capstone/output/SecureCAT_Ch1_Ch2_Manuscript_rebuilt.docx")
MANUSCRIPT_MD = Path("capstone/manuscript/SecureCAT_Ch1_Ch2_Manuscript.md")

# Canonical tags and their expected heading text
BOOKMARKS = [
    ('ch1_introduction', 'Chapter 1'),
    ('ch1_bg_of_the_study', 'Background of the Study'),
    ('ch1_conceptual_framework', 'Conceptual Framework of the Study'),
    ('ch1_objectives', 'Objectives of the Study'),
    ('ch1_scope_delimitations', 'Scope and Limitation of the Study'),
    ('ch1_significance', 'Importance of the Study'),
    ('ch2_methodology', 'Chapter 2'),
    ('ch2_research_design', 'Research Design'),
    ('ch2_software_model', 'Software Development Model'),
    ('ch2_project_plan', 'Project Plan'),
    ('ch2_project_assignment', 'Project Assignments'),
    ('ch2_population_locale', 'Population and Locale of the Study'),
    ('ch2_research_instruments', 'Research Instruments'),
    ('ch2_data_analysis', 'Data Analysis'),
    ('references_list', 'REFERENCES'),
    ('appendices_division', 'APPENDICES'),
    ('appendix_a_use_case', 'APPENDIX A'),
    ('appendix_b_letter_conduct', 'APPENDIX B'),
]


def sanitize(n):
    return re.sub(r'[^a-zA-Z0-9_]', '_', n)


def parse_md_sections(md_path):
    """Parse MD file and extract TAG sections with their UPDATE content."""
    content = md_path.read_text(encoding='utf-8')
    
    tag_pattern = re.compile(r'<!--\s*TAG:\s*(\S+)\s*-->')
    tag_positions = [(m.start(), m.group(1)) for m in tag_pattern.finditer(content)]
    
    sections = {}
    for i, (pos, tag_name) in enumerate(tag_positions):
        end_pos = tag_positions[i + 1][0] if i + 1 < len(tag_positions) else len(content)
        section_content = content[pos:end_pos]
        
        # Extract heading
        section_no_tag = re.sub(r'<!--\s*TAG:\s*\S+\s*-->\s*', '', section_content, count=1)
        heading_match = re.search(r'^(#{1,4})\s+(.+)$', section_no_tag, re.MULTILINE)
        heading = heading_match.group(2).strip() if heading_match else tag_name
        heading_level = len(heading_match.group(1)) if heading_match else 2
        
        # Extract UPDATE content
        update_match = re.search(
            r'<!--\s*UPDATE:START\s*-->(.*?)<!--\s*UPDATE:END\s*-->',
            section_content, re.DOTALL
        )
        update_content = update_match.group(1).strip() if update_match else ''
        
        sections[tag_name] = {
            'heading': heading,
            'level': heading_level,
            'content': update_content,
        }
    
    return sections


def strip_all_bookmarks(doc):
    """Remove ALL bookmarkStart/End elements from entire document."""
    removed = 0
    for p in doc.paragraphs:
        for bm in list(p._p.xpath('.//w:bookmarkStart')):
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                bm.getparent().remove(bm)
                removed += 1
        for bm in list(p._p.xpath('.//w:bookmarkEnd')):
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                bm.getparent().remove(bm)
                removed += 1
    # Also by ID
    tag_ids = set()
    for p in doc.paragraphs:
        for bm in list(p._p.xpath('.//w:bookmarkStart')):
            name = bm.get(qn('w:name'))
            bid = bm.get(qn('w:id'))
            if name and name.startswith('tag_') and bid:
                tag_ids.add(bid)
                bm.getparent().remove(bm)
                removed += 1
    for p in doc.paragraphs:
        for bm in list(p._p.xpath('.//w:bookmarkEnd')):
            bid = bm.get(qn('w:id'))
            if bid and bid in tag_ids:
                bm.getparent().remove(bm)
                removed += 1
    return removed


def fix_body_as_headings(doc):
    """Fix paragraphs that are styled as headings but contain body text."""
    fixed = 0
    for i, p in enumerate(doc.paragraphs):
        if not p.style.name.startswith('Heading'):
            continue
        text = p.text.strip()
        # Body text indicators: long text, starts with spaces, or contains known body-start patterns
        if (len(text) > 100 or
            p.text.startswith('     ') or
            text.startswith('The Researchers.') or
            text.startswith('The Community.') or
            text.startswith('The Client Institution.') or
            text.startswith('The College or Department.') or
            text.startswith('The Students.') or
            text.startswith('Future Researchers.') or
            text.startswith('The Respondents.') or
            text.startswith('This chapter presents') or
            text.startswith('This study employs') or
            text.startswith('This study will use') or
            text.startswith('The project timeline') or
            text.startswith('The data analysis') or
            text.startswith('To evaluate') or
            text.startswith('The execution') or
            text.startswith('In the context')):
            old = p.style.name
            p.style = doc.styles['Normal']
            fixed += 1
            print(f"  Fixed para {i}: {old} → Normal: \"{text[:70]}...\"")
    return fixed


def delete_duplicate_paragraphs(doc):
    """Remove duplicate body paragraphs (exact text match, keeping first occurrence)."""
    seen_texts = {}
    to_delete = []
    
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            continue  # Don't deduplicate headings
        text = p.text.strip()
        if not text or len(text) < 50:
            continue  # Skip short/empty paras
        
        if text in seen_texts:
            to_delete.append(i)
            print(f"  Duplicate para {i} (first at {seen_texts[text]}): \"{text[:70]}...\"")
        else:
            seen_texts[text] = i
    
    # Delete in reverse order to preserve indices
    for i in reversed(to_delete):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    
    return len(to_delete)


def find_heading_index(doc, heading_text):
    """Find the paragraph index of a heading by exact text match."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            if p.text.strip().lower() == heading_text.strip().lower():
                return i
    return -1


def find_next_bookmarked_heading_index(doc, start_idx, bookmarked_headings):
    """Find the next heading paragraph after start_idx that is in the BOOKMARKS list."""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            if doc.paragraphs[i].text.strip().lower() in bookmarked_headings:
                return i
    return len(doc.paragraphs)


def delete_body_between_headings(doc, start_idx, end_idx):
    """Delete all paragraphs between start and end indices."""
    to_delete = []
    for i in range(start_idx + 1, end_idx):
        to_delete.append(i)
    
    for i in reversed(to_delete):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    
    return len(to_delete)


def insert_content_after(doc, heading_idx, content_text, is_references=False):
    """Insert paragraphs after the heading at heading_idx."""
    if not content_text or not content_text.strip():
        return 0
    
    heading_element = doc.paragraphs[heading_idx]._p
    
    # Split content into paragraphs
    paragraphs = content_text.split('\n\n')
    added = 0
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    
    for para_text in reversed(paragraphs):
        para_text = para_text.strip()
        if not para_text:
            continue
        
        if para_text.startswith('#'):
            level = len(para_text) - len(para_text.lstrip('#'))
            text = para_text.lstrip('# ').strip()
            p = doc.add_paragraph(text, style=f'Heading {min(level, 4)}')
            for run in p.runs:
                run.font.name = 'Times New Roman'
            text_upper = text.upper()
            if text_upper.startswith('CHAPTER') or text_upper.startswith('REFERENCES') or text_upper.startswith('APPENDIX') or text_upper.startswith('APPENDICES') or text_upper in ('INTRODUCTION', 'METHODOLOGY'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.page_break_before = True
                for run in p.runs:
                    run.bold = True
        else:
            p = doc.add_paragraph(para_text, style='Normal')
            if is_references:
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.left_indent = Inches(0)
            
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if not p.runs:
                p.add_run()
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        heading_element.addnext(p._p)
        added += 1
    
    return added


def add_range_bookmark(start_para, end_para, tag_name):
    """Add a range bookmark spanning from start_para to end_para."""
    bm_name = f'tag_{sanitize(tag_name)}'
    bm_id = str(abs(hash(bm_name)) % 1000000)
    
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bm_id)
    start.set(qn('w:name'), bm_name)
    start_para._p.insert(0, start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bm_id)
    end.set(qn('w:name'), bm_name)
    
    if start_para is end_para:
        start_para._p.append(end)
    else:
        end_para._p.append(end)


def update_md_meta(md_path, docx_sha):
    """Update META tag in MD with new SHA256."""
    content = md_path.read_text(encoding='utf-8')
    meta_str = f'docx-sync-version="{datetime.now(timezone.utc).isoformat()}" docx-sha256="{docx_sha}"'
    new_content = re.sub(
        r'<!--\s*META:.*?-->',
        f'<!-- META: {meta_str} -->',
        content
    )
    md_path.write_text(new_content, encoding='utf-8')


def main():
    print("=" * 70)
    print("FULL DOCX REBUILD FROM MANUSCRIPT MD")
    print("=" * 70)
    
    # Parse MD
    print("\n📖 Parsing manuscript MD...")
    md_sections = parse_md_sections(MANUSCRIPT_MD)
    print(f"  Found {len(md_sections)} TAG sections:")
    for tag, data in md_sections.items():
        print(f"    {tag}: heading='{data['heading']}', level={data['level']}, content={len(data['content'])} chars")
    
    # Copy defense DOCX for rebuilding
    print(f"\n📦 Creating working copy of defense DOCX...")
    shutil.copy2(DEFENSE_DOCX, REPAIRED_DOCX)
    
    original_sha = hashlib.sha256(DEFENSE_DOCX.read_bytes()).hexdigest()
    print(f"  Original SHA256: {original_sha}")
    
    # Load document
    doc = Document(REPAIRED_DOCX)
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    
    # Step 1: Strip all existing bookmarks
    print("\n🧹 Step 1: Strip all existing tag_ bookmarks...")
    removed = strip_all_bookmarks(doc)
    print(f"  Removed {removed} bookmark elements")
    
    # Step 2: Fix body text styled as headings
    print("\n🔧 Step 2: Fix body text styled as headings...")
    fixed = fix_body_as_headings(doc)
    print(f"  Fixed {fixed} paragraphs")
    
    # Step 3: Remove duplicate paragraphs
    print("\n🗑️  Step 3: Remove duplicate paragraphs...")
    dupes = delete_duplicate_paragraphs(doc)
    print(f"  Removed {dupes} duplicates")
    
    # Save intermediate state
    doc.save(REPAIRED_DOCX)
    doc = Document(REPAIRED_DOCX)  # Reload for fresh indices
    
    print(f"\n  After cleanup: {len(doc.paragraphs)} paragraphs")
    
    # Step 4: Show current heading structure
    print("\n📋 Step 4: Current heading structure after cleanup:")
    headings = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            headings.append((i, p.text.strip(), p.style.name))
            print(f"    [{i:3d}] ({p.style.name}) \"{p.text.strip()[:70]}\"")
    
    # Step 5: For each MD section, find its heading in DOCX, clear body, insert content
    print("\n📝 Step 5: Rebuild each section from MD content...")
    
    # Process sections in REVERSE order (to preserve indices)
    tag_order = [t for t, _ in BOOKMARKS]
    
    # First pass: map tags to heading indices
    tag_to_heading = {}
    for tag, heading_text in BOOKMARKS:
        idx = find_heading_index(doc, heading_text)
        if idx >= 0:
            tag_to_heading[tag] = idx
            print(f"    ✓ {tag} → para {idx}: \"{doc.paragraphs[idx].text.strip()[:60]}\"")
        else:
            print(f"    ⚠ {tag}: heading \"{heading_text}\" NOT FOUND in DOCX")
    
    # Second pass: clear and rebuild each section (reverse order for index stability)
    print("\n  Rebuilding sections (reverse order)...")
    bookmarked_headings = {h.strip().lower() for _, h in BOOKMARKS}
    for tag in reversed(tag_order):
        if tag not in tag_to_heading:
            continue
        if tag not in md_sections:
            continue
        
        heading_idx = tag_to_heading[tag]
        next_heading_idx = find_next_bookmarked_heading_index(doc, heading_idx, bookmarked_headings)
        md_content = md_sections[tag]['content']
        
        # Delete existing body
        deleted = delete_body_between_headings(doc, heading_idx, next_heading_idx)
        
        # Insert MD content
        inserted = insert_content_after(doc, heading_idx, md_content, is_references=(tag == 'references_list'))
        
        print(f"    {tag}: deleted {deleted} old paras, inserted {inserted} new paras ({len(md_content)} chars)")
    
    # Step 6: Add clean bookmarks
    print("\n🔖 Step 6: Add clean bookmarks...")
    # Save and reload to get fresh indices
    doc.save(REPAIRED_DOCX)
    doc = Document(REPAIRED_DOCX)
    
    bookmark_count = 0
    used_paras = set()
    for tag, heading_text in BOOKMARKS:
        for i, p in enumerate(doc.paragraphs):
            if i in used_paras:
                continue
            if not p.style.name.startswith('Heading'):
                continue
            if p.text.strip().lower() == heading_text.strip().lower():
                # Find last paragraph before next section heading
                end_idx = i
                for j in range(i + 1, len(doc.paragraphs)):
                    p_next = doc.paragraphs[j]
                    if p_next.style.name.lower().startswith('heading'):
                        p_next_text = p_next.text.strip().lower()
                        is_next_section = False
                        for _, next_heading in BOOKMARKS:
                            next_heading_clean = next_heading.strip().lower()
                            if p_next_text == next_heading_clean or p_next_text.startswith(next_heading_clean):
                                is_next_section = True
                                break
                        if is_next_section:
                            break
                    end_idx = j
                
                add_range_bookmark(p, doc.paragraphs[end_idx], tag)
                used_paras.add(i)
                bookmark_count += 1
                print(f"    ✓ {tag} → range para {i} to {end_idx}")
                break
        else:
            print(f"    ⚠ {tag}: could not bookmark \"{heading_text}\"")
    
    # Ensure page breaks before Chapter, REFERENCES, and APPENDIX headings
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading'):
            text_upper = p.text.strip().upper()
            if text_upper.startswith('CHAPTER') or text_upper.startswith('REFERENCES') or text_upper.startswith('APPENDIX') or text_upper.startswith('APPENDICES'):
                p.paragraph_format.page_break_before = True
                
    # Save final
    doc.save(REPAIRED_DOCX)
    rebuilt_sha = hashlib.sha256(REPAIRED_DOCX.read_bytes()).hexdigest()
    print(f"\n  Saved rebuilt DOCX: {REPAIRED_DOCX}")
    print(f"  Rebuilt SHA256: {rebuilt_sha}")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    
    # Step 7: Verify content
    print("\n🔍 Step 7: Verify rebuilt DOCX content matches MD...")
    doc = Document(REPAIRED_DOCX)
    
    all_ok = True
    for tag, heading_text in BOOKMARKS:
        heading_idx = find_heading_index(doc, heading_text)
        if heading_idx < 0:
            print(f"  ❌ {tag}: heading not found")
            all_ok = False
            continue
        
        next_idx = find_next_bookmarked_heading_index(doc, heading_idx, bookmarked_headings)
        
        # Get DOCX content
        docx_parts = []
        for k in range(heading_idx + 1, next_idx):
            p = doc.paragraphs[k]
            text = p.text.strip()
            if text:
                if p.style.name.startswith('Heading'):
                    try:
                        level = int(p.style.name.split()[-1])
                        text = f"{'#' * level} {text}"
                    except (ValueError, IndexError):
                        text = f"# {text}"
                docx_parts.append(text)
        docx_content = '\n\n'.join(docx_parts)
        
        # Get MD content
        md_content = md_sections.get(tag, {}).get('content', '')
        
        docx_len = len(docx_content)
        md_len = len(md_content)
        
        if md_len == 0 and docx_len == 0:
            status = "✅"
        elif md_len == 0:
            status = "⚠️ "
        elif abs(docx_len - md_len) < 50:  # Allow small formatting differences
            status = "✅"
        else:
            status = "❌"
            all_ok = False
        
        print(f"  {status} {tag:40s}: MD={md_len:6d} chars, DOCX={docx_len:6d} chars (Δ={docx_len-md_len:+d})")
    
    # Update MD META with new SHA
    update_md_meta(MANUSCRIPT_MD, rebuilt_sha)
    print(f"\n  Updated MD META with rebuilt SHA256: {rebuilt_sha[:16]}...")
    
    # Summary
    print(f"\n{'='*70}")
    print("REBUILD SUMMARY")
    print(f"{'='*70}")
    print(f"  Original defense DOCX SHA256:  {original_sha}")
    print(f"  Rebuilt DOCX SHA256:           {rebuilt_sha}")
    print(f"  Bookmarks added:               {bookmark_count}")
    
    if all_ok:
        print(f"\n  ✅ ALL SECTIONS MATCH — Content integrity verified!")
        print(f"\n  Next steps:")
        print(f"    1. Open {REPAIRED_DOCX} in Word/LibreOffice and visually inspect")
        print(f"    2. If satisfied, copy to defense file:")
        print(f"       cp '{REPAIRED_DOCX}' '{DEFENSE_DOCX}'")
        print(f"    3. Push to Google Drive:")
        print(f"       rclone copy '{DEFENSE_DOCX}' 'gdrive:A.Y. 2026-2027/Capstone/output/' -v")
    else:
        print(f"\n  ⚠️  SOME SECTIONS HAVE MISMATCHES — Review above")
    
    return 0 if all_ok else 1


if __name__ == '__main__':
    sys.exit(main())

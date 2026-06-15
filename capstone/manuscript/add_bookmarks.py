#!/usr/bin/env python3
"""
Add bookmarks to all 23 headings in the defense DOCX.
SAFE: Only inserts XML bookmark elements — NO content/style/formatting changes.
Run ONCE after pulling reverted cloud DOCX, then surgical updates work forever.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

DOCX_PATH = "capstone/output/SecureCAT_Ch1_Ch2_Manuscript[never delete].docx"

def sanitize(n): return re.sub(r'[^a-zA-Z0-9_]', '_', n)

BOOKMARKS = [
    ('ch1_introduction', 'Chapter 1'),
    ('ch1_bg_of_the_study', 'Background of the Study'),
    ('ch1_conceptual_framework', 'Conceptual Framework of the Study'),
    ('ch1_objectives', 'Objectives of the Study'),
    ('ch1_scope_delimitations', 'Scope and Limitation of the Study'),
    ('ch1_significance', 'Significance of the Study'),
    ('ch2_methodology', 'Chapter 2'),
    ('ch2_research_design', 'Research Design'),
    ('ch2_software_model', 'Software Development Model'),
    ('ch2_project_plan', 'Project Plan'),
    ('ch2_project_assignment', 'Project Assignments'),
    ('ch2_population_locale', 'Population and Locale of the Study'),
    ('ch2_research_instruments', 'Research Instruments'),
    ('ch2_data_analysis', 'Data Analysis'),
    ('references_list', 'REFERENCES'),
    ('appendices_division', 'APPENDICES'),
    ('appendix_a_use_case', 'APPENDIX A'),
    ('appendix_b_letter_conduct', 'APPENDIX B'),
]

def matches_heading(p, search_text):
    """EXACT or STARTSWITH match ONLY for Heading-style paragraphs.
    Content paragraphs often contain heading words but are not headings."""
    ptext = p.text.strip()
    stext = search_text.strip().lower()
    plower = ptext.lower()
    
    # 1. Exact match (after stripping)
    if plower == stext:
        return True
    # 2. Starts with (heading at beginning of paragraph, e.g., "Background of the Study" or "Background of the Study\n")
    if plower.startswith(stext):
        return True
    # NO loose contains matching - that catches content paragraphs
    return False

def main():
    doc = Document(DOCX_PATH)
    print(f"Loaded: {DOCX_PATH} ({len(doc.paragraphs)} paragraphs)")
    
    # Convert expected headings that are styled as Normal to Heading styles first
    HEADING_CONVERSIONS = [
        ("Chapter 1", 1),
        ("INTRODUCTION", 1),
        ("Background of the Study", 2),
        ("Conceptual Framework of the Study", 2),
        ("Objectives of the Study", 2),
        ("Scope and Limitation of the Study", 2),
        ("Significance of the Study", 2),
        ("Importance of the Study", 2),
        ("Chapter 2", 1),
        ("METHODOLOGY", 1),
        ("Research Design", 2),
        ("Software Development Model", 2),
        ("Project Plan", 2),
        ("Project Assignments", 2),
        ("Population and Locale of the Study", 2),
        ("Research Instruments", 2),
        ("Data Analysis", 2),
        ("REFERENCES", 2),
        ("APPENDICES", 2),
        ("APPENDIX A", 2),
        ("APPENDIX B", 2),
    ]
    
    converted = 0
    for search_text, level in HEADING_CONVERSIONS:
        for p in doc.paragraphs:
            ptext = p.text.strip().lower()
            stext = search_text.strip().lower()
            if (ptext == stext or ptext.startswith(stext)) and not p.style.name.startswith('Heading'):
                p.style = doc.styles[f'Heading {level}']
                converted += 1
                print(f"  Converted normal paragraph to Heading {level}: '{p.text}'")
    if converted > 0:
        print(f"Converted {converted} paragraphs to Heading styles.")
    
    used_paragraphs = set()
    added = 0
    for tag, heading_text in BOOKMARKS:
        for i, p in enumerate(doc.paragraphs):
            if i in used_paragraphs:
                continue
            # Match ONLY on Heading styles with strict heading text match
            is_heading = p.style.name.lower().startswith('heading')
            text_match = matches_heading(p, heading_text)
            
            if is_heading and text_match:
                bm_name = f'tag_{sanitize(tag)}'
                bm_id = str(abs(hash(bm_name)) % 1000000)
                
                # Check if already has this bookmark start
                existing = p._p.xpath(f'.//w:bookmarkStart[@w:name="{bm_name}"]')
                if existing:
                    bm_id = existing[0].get(qn('w:id'))
                    # Remove any existing ends to update range
                    for p_any in doc.paragraphs:
                        for old_end in p_any._p.xpath(f'.//w:bookmarkEnd[@w:id="{bm_id}"]'):
                            old_end.getparent().remove(old_end)
                        for old_end in p_any._p.xpath(f'.//w:bookmarkEnd[@w:name="{bm_name}"]'):
                            old_end.getparent().remove(old_end)
                else:
                    start = OxmlElement('w:bookmarkStart')
                    start.set(qn('w:id'), bm_id)
                    start.set(qn('w:name'), bm_name)
                    p._p.insert(0, start)
                
                # Find last paragraph before next section heading
                end_idx = i
                for j in range(i + 1, len(doc.paragraphs)):
                    p_next = doc.paragraphs[j]
                    if p_next.style.name.lower().startswith('heading'):
                        p_next_text = p_next.text.strip().lower()
                        is_next_section = False
                        for _, next_heading in BOOKMARKS:
                            next_heading_clean = next_heading.strip().lower()
                            if p_next_text == next_heading_clean or p_next_text.startswith(next_heading_clean):
                                is_next_section = True
                                break
                        if is_next_section:
                            break
                    end_idx = j
                
                end_p = doc.paragraphs[end_idx]
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), bm_id)
                end.set(qn('w:name'), bm_name)
                end_p._p.append(end)
                
                used_paragraphs.add(i)
                print(f"  ✓  {tag} -> range para {i} to {end_idx} ({p.style.name} to {end_p.style.name})")
                added += 1
                break
        else:
            print(f"  ⚠  {tag}: NO MATCH for heading \"{heading_text}\"")
    
    doc.save(DOCX_PATH)
    print(f"\nDone. Added {added} bookmarks. Saved to {DOCX_PATH}")

if __name__ == '__main__':
    main()
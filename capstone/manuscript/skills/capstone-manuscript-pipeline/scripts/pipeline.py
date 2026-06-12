#!/usr/bin/env python3
"""
Capstone Manuscript Pipeline — Unified workflow with guardrails.
Orchestrates: pull → drift check → surgical update → push → backup.
"""
import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from docx import Document

# ======== CONFIG ========
DEFENSE_FILE = "SecureCAT_Ch1_Ch2_Manuscriptnever_delete.docx"
DRIVE_BASE = "gdrive:A.Y. 2026-2027/Capstone/output/"
MANUSCRIPT_MD = Path("capstone/manuscript/SecureCAT_Ch1_Ch2_Manuscript.md")
TEMPLATE_DOCX = Path("capstone/templates/BSIT Capstone Template.docx")
OUTPUT_DIR = Path("capstone/output")
BACKUP_DIR = OUTPUT_DIR / "backups"
LOCAL_DOCX = OUTPUT_DIR / DEFENSE_FILE
REMOTE_DOCX = f"{DRIVE_BASE}{DEFENSE_FILE}"

MAX_BACKUPS = 30

# ======== UTILITIES ========

def run_cmd(cmd, check=True, capture=True):
    """Run shell command, return (exit_code, stdout, stderr)."""
    try:
        result = subprocess.run(cmd, shell=True, capture_output=capture, text=True, timeout=120)
        return result.returncode, result.stdout.strip(), result.stderr.strip()
    except subprocess.TimeoutExpired:
        return -1, "", "Command timed out"
    except Exception as e:
        return -1, "", str(e)

def compute_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()

def parse_meta(md_path: Path) -> dict:
    content = md_path.read_text(encoding='utf-8')
    match = re.search(r'<!--\s*META:\s*(.*?)\s*-->', content)
    if not match:
        return {}
    meta = {}
    for part in match.group(1).split():
        if '=' in part:
            k, v = part.split('=', 1)
            meta[k] = v.strip('"')
    return meta

def update_meta(md_path: Path, new_meta: dict):
    content = md_path.read_text(encoding='utf-8')
    meta_str = ' '.join(f'{k}="{v}"' for k, v in new_meta.items())
    new_content = re.sub(
        r'<!--\s*META:.*?-->',
        f'<!-- META: {meta_str} -->',
        content
    )
    if '<!-- META:' not in content:
        new_content = f'<!-- META: {meta_str} -->\n\n{new_content}'
    md_path.write_text(new_content, encoding='utf-8')

def ensure_backup_dir():
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)

def create_backup():
    """Create timestamped backup of current defense file."""
    ensure_backup_dir()
    if not LOCAL_DOCX.exists():
        return None
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"SecureCAT_Ch1_Ch2_Manuscript_backup_{timestamp}.docx"
    backup_path = BACKUP_DIR / backup_name
    shutil.copy2(LOCAL_DOCX, backup_path)
    print(f"📦 Backup created: {backup_path}")
    
    # Prune old backups
    backups = sorted(BACKUP_DIR.glob("*.docx"))
    for old in backups[:-MAX_BACKUPS]:
        old.unlink()
        print(f"🗑️  Pruned old backup: {old.name}")
    return backup_path

# ======== ANTI-DELETION GUARD ========

def anti_deletion_guard() -> bool:
    """Verify defense file exists on Drive and locally."""
    print("🛡️  Anti-deletion guard: checking defense file...")
    
    # Check Drive
    code, stdout, stderr = run_cmd(f'rclone ls "{REMOTE_DOCX}"')
    if code != 0 or not stdout.strip():
        print(f"❌ DEFENSE FILE MISSING ON DRIVE: {REMOTE_DOCX}")
        print("   Cannot proceed — file may have been deleted.")
        return False
    
    # Check local
    if not LOCAL_DOCX.exists():
        print(f"⚠️  Local copy missing: {LOCAL_DOCX}")
        print("   Will pull from Drive...")
    else:
        # Verify SHA256 matches
        local_sha = compute_sha256(LOCAL_DOCX)
        # Can't easily get remote SHA without downloading, skip for speed
        print(f"✅ Local defense file exists: {local_sha[:16]}...")
    
    print("✅ Anti-deletion guard passed")
    return True

# ======== DRIFT CHECK ========

def drift_check() -> tuple[int, dict]:
    """Check drift between local DOCX and manuscript MD META.
    Returns (status_code, meta_dict).
    Status: 0=PASS, 1=DRIFT, 2=MISSING_LOCAL, 3=NO_META
    """
    if not LOCAL_DOCX.exists():
        print("❌ Local defense file not found for drift check")
        return 2, {}
    
    meta = parse_meta(MANUSCRIPT_MD)
    expected_sha = meta.get('docx-sha256')
    sync_version = meta.get('docx-sync-version', 'unknown')
    
    if not expected_sha:
        print("❌ NO_META: No docx-sha256 in manuscript MD META tag")
        print("   Recommend: run full rebuild or accept-cloud to initialize")
        return 3, meta
    
    actual_sha = compute_sha256(LOCAL_DOCX)
    
    if actual_sha == expected_sha:
        print(f"✅ PASS: DOCX matches META (synced {sync_version})")
        return 0, meta
    
    # DRIFT DETECTED
    print(f"❌ DRIFT_DETECTED: DOCX has unwatched changes")
    print(f"   Expected SHA256: {expected_sha} (from MD META, synced {sync_version})")
    print(f"   Actual SHA256:   {actual_sha}")
    
    current_text = extract_docx_text(LOCAL_DOCX)
    print(f"   Current DOCX length: {len(current_text)} chars")
    return 1, meta

def extract_docx_text(path: Path) -> str:
    try:
        doc = Document(path)
        return '\n'.join(p.text for p in doc.paragraphs)
    except:
        return ""

# ======== DRIFT RESOLUTION ========

def resolve_drift(meta: dict) -> bool:
    """Interactive drift resolution. Returns True if resolved, False if aborted."""
    print("""
Resolution options:
  [1] accept-cloud   → Use cloud version, update META, continue pipeline
  [2] keep-local     → Push local DOCX to cloud (overwrite cloud)
  [3] merge          → Manual merge required (open both, resolve, re-run)
  [4] snapshot-cloud → Archive cloud version to backups/, then proceed
  [5] abort          → Cancel pipeline
""")
    while True:
        choice = input("Enter choice [1-5]: ").strip()
        if choice == '1':
            return accept_cloud_resolution(meta)
        elif choice == '2':
            return keep_local_resolution()
        elif choice == '3':
            merge_resolution()
            return False  # Need re-run after manual merge
        elif choice == '4':
            return snapshot_cloud_resolution()
        elif choice == '5':
            print("🛑 Pipeline aborted by user")
            return False
        else:
            print("Invalid choice. Enter 1-5.")

def accept_cloud_resolution(meta: dict) -> bool:
    """Accept cloud version as truth: update META tag with new SHA256."""
    print("☁️  Accepting cloud version...")
    actual_sha = compute_sha256(LOCAL_DOCX)
    new_meta = {
        'docx-sync-version': datetime.now(timezone.utc).isoformat(),
        'docx-sha256': actual_sha
    }
    update_meta(MANUSCRIPT_MD, new_meta)
    print(f"✅ META updated: SHA256={actual_sha[:16]}...")
    return True

def keep_local_resolution() -> bool:
    """Push local DOCX to cloud, overwriting cloud version."""
    print("📤 Keeping local — pushing to Drive...")
    code, stdout, stderr = run_cmd(f'rclone copy "{LOCAL_DOCX}" "{DRIVE_BASE}" --progress -v')
    if code != 0:
        print(f"❌ Push failed: {stderr}")
        return False
    print("✅ Local version pushed to Drive")
    return True

def merge_resolution():
    """Guide user through manual merge."""
    print("""
📝 MANUAL MERGE REQUIRED:

1. Open both versions:
   - Cloud: Download from Drive or open in Google Docs
   - Local: {LOCAL_DOCX}

2. Merge changes manually in Word/Google Docs

3. Save merged version to:
   {LOCAL_DOCX}

4. Re-run pipeline:
   python pipeline.py --action full
""")

def snapshot_cloud_resolution() -> bool:
    """Archive cloud version, then proceed."""
    ensure_backup_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    snapshot_name = f"cloud_snapshot_{timestamp}.docx"
    snapshot_path = BACKUP_DIR / snapshot_name
    
    print(f"📸 Snapshotting cloud version to {snapshot_path}...")
    code, stdout, stderr = run_cmd(f'rclone copy "{REMOTE_DOCX}" "{BACKUP_DIR}" --progress -v')
    if code != 0:
        print(f"❌ Snapshot failed: {stderr}")
        return False
    
    # Rename downloaded file to snapshot name
    downloaded = BACKUP_DIR / DEFENSE_FILE
    if downloaded.exists():
        downloaded.rename(snapshot_path)
        print(f"✅ Cloud snapshot saved: {snapshot_path}")
    
    # Update local with cloud version (already downloaded)
    # Now accept cloud via META update
    meta = parse_meta(MANUSCRIPT_MD)
    return accept_cloud_resolution(meta)

# ======== PULL FROM DRIVE ========

def pull_from_drive() -> bool:
    """Pull latest defense file from Drive to local."""
    print(f"⬇️  Pulling {DEFENSE_FILE} from Drive...")
    code, stdout, stderr = run_cmd(f'rclone copy "{REMOTE_DOCX}" "{OUTPUT_DIR}" --progress -v')
    if code != 0:
        print(f"❌ Pull failed: {stderr}")
        return False
    print("✅ Pull complete")
    return True

# ======== SURGICAL UPDATE ========

def surgical_update() -> bool:
    """Run md-updates-docx surgical update."""
    print("🔧 Running surgical MD→DOCX update...")
    
    # Import and run md_to_docx functions
    sys.path.insert(0, str(Path("capstone/manuscript")))
    from md_to_docx import parse_manuscript, apply_updates, compute_docx_sha256, update_manuscript_meta
    from docx import Document
    
    meta, tags = parse_manuscript(MANUSCRIPT_MD)
    doc = Document(LOCAL_DOCX)
    
    # Build tag map
    from md_to_docx import build_tag_map
    tag_map = build_tag_map(doc, tags)
    
    # Apply all tags with UPDATE/REMOVE
    update_tags = list(tags.keys())
    apply_updates(doc, tags, tag_map, update_tags)
    
    # Save
    doc.save(LOCAL_DOCX)
    
    # Update META
    new_sha = compute_docx_sha256(LOCAL_DOCX)
    new_meta = {
        'docx-sync-version': datetime.now(timezone.utc).isoformat(),
        'docx-sha256': new_sha
    }
    update_manuscript_meta(MANUSCRIPT_MD, new_meta)
    print(f"✅ Surgical update complete. New SHA256: {new_sha[:16]}...")
    return True

# ======== PUSH TO DRIVE ========

def push_to_drive() -> bool:
    """Push updated defense file to Drive (additive copy only)."""
    print(f"⬆️  Pushing {DEFENSE_FILE} to Drive...")
    code, stdout, stderr = run_cmd(f'rclone copy "{LOCAL_DOCX}" "{DRIVE_BASE}" --progress -v')
    if code != 0:
        print(f"❌ Push failed: {stderr}")
        return False
    print("✅ Push complete")
    return True

# ======== MAIN ACTIONS ========

def action_pull():
    if not anti_deletion_guard():
        return False
    return pull_from_drive()

def action_check():
    if not anti_deletion_guard():
        return False
    if not pull_from_drive():
        return False
    status, meta = drift_check()
    return status == 0

def action_update():
    if not anti_deletion_guard():
        return False
    if not pull_from_drive():
        return False
    
    status, meta = drift_check()
    if status == 1:  # DRIFT
        if not resolve_drift(meta):
            return False
        # Re-check after resolution
        status, meta = drift_check()
        if status != 0:
            print("❌ Drift not resolved")
            return False
    elif status != 0:
        print(f"❌ Drift check failed (code {status})")
        return False
    
    # No drift — run surgical update
    return surgical_update()

def action_push():
    if not anti_deletion_guard():
        return False
    return push_to_drive()

def action_full():
    """Full pipeline: pull → check → update → push → backup"""
    print("🚀 CAPSTONE MANUSCRIPT PIPELINE — FULL RUN")
    print("=" * 50)
    
    # 1. Anti-deletion guard
    if not anti_deletion_guard():
        return False
    
    # 2. Pull latest from Drive
    if not pull_from_drive():
        return False
    
    # 3. Drift check + resolution
    status, meta = drift_check()
    if status == 1:  # DRIFT
        print("\n" + "=" * 50)
        if not resolve_drift(meta):
            return False
        # Re-check after resolution
        status, meta = drift_check()
        if status != 0:
            print("❌ Drift not resolved after resolution")
            return False
    elif status != 0:
        print(f"❌ Drift check failed (code {status})")
        return False
    
    # 4. Surgical update (if no drift, or after resolution)
    if not surgical_update():
        return False
    
    # 5. Push to Drive
    if not push_to_drive():
        return False
    
    # 6. Create local backup
    create_backup()
    
    print("\n" + "=" * 50)
    print("🎉 PIPELINE COMPLETE — All guards passed")
    return True

# ======== CLI ========

def main():
    parser = argparse.ArgumentParser(description="Capstone Manuscript Pipeline")
    parser.add_argument('--action', choices=['pull', 'check', 'update', 'push', 'full'],
                       default='full', help="Pipeline action to run")
    args = parser.parse_args()
    
    action_map = {
        'pull': action_pull,
        'check': action_check,
        'update': action_update,
        'push': action_push,
        'full': action_full,
    }
    
    success = action_map[args.action]()
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()

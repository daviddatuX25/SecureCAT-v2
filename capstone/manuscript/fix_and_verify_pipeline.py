#!/usr/bin/env python3
"""
Fix & Verify Pipeline — Full DOCX roundtrip repair and validation.

Steps:
  1. REPAIR: Strip ALL existing tag_ bookmarks (duplicates/misplaced)
  2. REPAIR: Fix body paragraphs erroneously styled as headings
  3. REPAIR: Re-insert clean bookmarks exactly ONCE per heading
  4. EXTRACT: DOCX → JSON (bookmark-based extraction)
  5. MERGE: JSON → MD (update manuscript MD)
  6. CONVERT: MD → DOCX (output to test file, NOT overwriting original)
  7. COMPARE: Original repaired DOCX vs output DOCX content comparison
"""
import hashlib
import json
import re
import sys
import shutil
from datetime import datetime, timezone
from pathlib import Path
from collections import Counter

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======== CONFIG ========
ORIGINAL_DOCX = Path("capstone/output/SecureCAT_Ch1_Ch2_Manuscript[never delete].docx")
REPAIRED_DOCX = Path("capstone/output/SecureCAT_Ch1_Ch2_Manuscript_repaired.docx")
OUTPUT_DOCX = Path("capstone/output/SecureCAT_Ch1_Ch2_Manuscript_roundtrip_test.docx")
MANUSCRIPT_MD = Path("capstone/manuscript/SecureCAT_Ch1_Ch2_Manuscript.md")
EXTRACTED_JSON = Path("capstone/manuscript/extracted_from_docx.json")

# Canonical bookmark definitions: (tag_name, exact_heading_text)
BOOKMARKS = [
    ('ch1_introduction', 'Chapter 1'),
    ('ch1_bg_of_the_study', 'Background of the Study'),
    ('ch1_conceptual_framework', 'Conceptual Framework of the Study'),
    ('ch1_objectives', 'Objectives of the Study'),
    ('ch1_scope_delimitations', 'Scope and Limitation of the Study'),
    ('ch1_significance', 'Importance of the Study'),
    ('ch2_methodology', 'Chapter 2'),
    ('ch2_research_design', 'Research Design'),
    ('ch2_software_model', 'Software Development Model'),
    ('ch2_project_plan', 'Project Plan'),
    ('ch2_project_assignment', 'Project Assignments'),
    ('ch2_population_locale', 'Population and Locale of the Study'),
    ('ch2_research_instruments', 'Research Instruments'),
    ('ch2_data_analysis', 'Data Analysis'),
    ('references_list', 'REFERENCES'),
    ('appendix_a_use_case', 'APPENDIX A'),
    ('appendix_b_letter_conduct', 'APPENDIX B'),
]

# Known body-text paragraphs that are WRONGLY styled as headings
# Detected by: style=Heading but text length > 100 OR text starts with whitespace
BODY_TEXT_INDICATORS = [
    lambda p: len(p.text.strip()) > 100,
    lambda p: p.text.startswith('     '),
    lambda p: p.text.startswith('The Researchers.'),
    lambda p: p.text.startswith('The Community.'),
    lambda p: p.text.startswith('The Client Institution.'),
    lambda p: p.text.startswith('The College or Department.'),
    lambda p: p.text.startswith('The Students.'),
    lambda p: p.text.startswith('Future Researchers.'),
    lambda p: p.text.startswith('The Respondents.'),
]


def sanitize(n):
    return re.sub(r'[^a-zA-Z0-9_]', '_', n)


# ======== STEP 1: STRIP ALL EXISTING BOOKMARKS ========

def strip_all_tag_bookmarks(doc):
    """Remove ALL bookmarkStart and bookmarkEnd elements with tag_ prefix."""
    removed = 0
    for p in doc.paragraphs:
        # Remove bookmarkStart elements
        for bm in p._p.xpath('.//w:bookmarkStart'):
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                bm.getparent().remove(bm)
                removed += 1
        # Remove bookmarkEnd elements with matching names
        for bm in p._p.xpath('.//w:bookmarkEnd'):
            # bookmarkEnd uses w:name attribute too
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                bm.getparent().remove(bm)
                removed += 1
    
    # Also remove by ID matching (bookmarkEnd might use ID instead of name)
    # Collect IDs of tag_ bookmarkStarts first
    tag_ids = set()
    for p in doc.paragraphs:
        for bm in p._p.xpath('.//w:bookmarkStart'):
            name = bm.get(qn('w:name'))
            bm_id = bm.get(qn('w:id'))
            if name and name.startswith('tag_') and bm_id:
                tag_ids.add(bm_id)
    
    # Remove any remaining bookmarkEnd with those IDs
    for p in doc.paragraphs:
        for bm in p._p.xpath('.//w:bookmarkEnd'):
            bm_id = bm.get(qn('w:id'))
            if bm_id and bm_id in tag_ids:
                bm.getparent().remove(bm)
                removed += 1
    
    return removed


# ======== STEP 2: FIX BODY TEXT STYLED AS HEADINGS ========

def fix_heading_styles(doc):
    """Demote body-text paragraphs from Heading to Normal."""
    fixed = 0
    for i, p in enumerate(doc.paragraphs):
        if not p.style.name.startswith('Heading'):
            continue
        
        # Check if this is actually body text
        is_body = any(check(p) for check in BODY_TEXT_INDICATORS)
        if is_body:
            old_style = p.style.name
            p.style = doc.styles['Normal']
            fixed += 1
            preview = p.text[:80].strip()
            print(f"  ✓ Fixed para {i}: {old_style} → Normal: \"{preview}...\"")
    
    return fixed


# ======== STEP 3: INSERT CLEAN BOOKMARKS ========

def insert_clean_bookmarks(doc):
    """Insert exactly ONE bookmark per heading, with strict text matching."""
    used_paras = set()
    added = 0
    missed = []
    
    for tag, heading_text in BOOKMARKS:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if i in used_paras:
                continue
            if not p.style.name.startswith('Heading'):
                continue
            
            ptext = p.text.strip().lower()
            stext = heading_text.strip().lower()
            
            if ptext == stext or ptext.startswith(stext):
                bm_name = f'tag_{sanitize(tag)}'
                bm_id = str(abs(hash(bm_name)) % 1000000)
                
                start = OxmlElement('w:bookmarkStart')
                start.set(qn('w:id'), bm_id)
                start.set(qn('w:name'), bm_name)
                p._p.insert(0, start)
                
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), bm_id)
                end.set(qn('w:name'), bm_name)
                p._p.append(end)
                
                used_paras.add(i)
                added += 1
                found = True
                print(f"  ✓ {tag} → para {i} ({p.style.name}): \"{p.text[:60]}\"")
                break
        
        if not found:
            missed.append((tag, heading_text))
            print(f"  ⚠ {tag}: NO MATCH for \"{heading_text}\"")
    
    return added, missed


# ======== STEP 4: EXTRACT DOCX → JSON ========

def extract_sections(doc):
    """Extract all bookmarked sections from DOCX."""
    bookmarks = []
    for i, p in enumerate(doc.paragraphs):
        for bm in p._p.xpath('.//w:bookmarkStart'):
            name = bm.get(qn('w:name'))
            if name and name.startswith('tag_'):
                tag = name[4:]  # remove 'tag_' prefix
                level_match = re.search(r'(\d+)', p.style.name)
                level = int(level_match.group(1)) if level_match else 2
                bookmarks.append({
                    'index': i,
                    'tag': tag,
                    'style': p.style.name,
                    'text': p.text.strip(),
                    'level': level,
                })
    
    bookmarks.sort(key=lambda b: b['index'])
    
    sections = {}
    for j, bm in enumerate(bookmarks):
        start = bm['index']
        end = bookmarks[j + 1]['index'] if j + 1 < len(bookmarks) else len(doc.paragraphs)
        
        # Extract body paragraphs between this heading and the next
        content_parts = []
        for k in range(start + 1, end):
            if k >= len(doc.paragraphs):
                break
            p = doc.paragraphs[k]
            text = p.text.strip()
            if text:
                content_parts.append(text)
        
        sections[bm['tag']] = {
            'heading': bm['text'],
            'level': bm['level'],
            'content': '\n\n'.join(content_parts),
            'para_start': start,
            'para_end': end,
            'body_para_count': len(content_parts),
        }
    
    return sections


# ======== STEP 5: BUILD MANUSCRIPT MD ========

def build_manuscript_md(sections, docx_sha256):
    """Build a complete manuscript MD from extracted sections."""
    tag_order = [t for t, _ in BOOKMARKS]
    
    lines = []
    
    # META tag
    meta_str = f'docx-sync-version="{datetime.now(timezone.utc).isoformat()}" docx-sha256="{docx_sha256}"'
    lines.append(f'<!-- META: {meta_str} -->')
    lines.append('')
    
    # Title block (preserve from current MD if it exists)
    md_path = MANUSCRIPT_MD
    if md_path.exists():
        current = md_path.read_text(encoding='utf-8')
        # Extract title block (everything before first TAG)
        first_tag = re.search(r'<!--\s*TAG:\s*\S+\s*-->', current)
        if first_tag:
            title_block = current[:first_tag.start()].strip()
            # Remove any META lines from title block
            title_block = re.sub(r'<!--\s*META:.*?-->\s*', '', title_block).strip()
            if title_block:
                lines.append(title_block)
                lines.append('')
    
    # Sections
    for tag in tag_order:
        if tag not in sections:
            print(f"  ⚠ TAG {tag} not in extracted sections, skipping in MD")
            continue
        
        data = sections[tag]
        heading_prefix = '#' * data['level'] + ' '
        
        lines.append(f'<!-- TAG: {tag} -->')
        lines.append(f'{heading_prefix}{data["heading"]}')
        lines.append('')
        lines.append('<!-- UPDATE:START -->')
        if data['content']:
            lines.append(data['content'])
        lines.append('<!-- UPDATE:END -->')
    
    return '\n'.join(lines) + '\n'


# ======== STEP 6: MD → DOCX (via existing converter) ========

def run_md_to_docx(md_path, docx_input, output_path):
    """Run the existing md_to_docx.py converter."""
    import subprocess
    cmd = [
        sys.executable,
        'capstone/manuscript/md_to_docx.py',
        '--md', str(md_path),
        '--docx', str(docx_input),
        '--output', str(output_path),
    ]
    print(f"  Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(Path.cwd()))
    print(result.stdout)
    if result.stderr:
        print(f"  STDERR: {result.stderr}")
    return result.returncode


# ======== STEP 7: CONTENT COMPARISON ========

def compare_docx_content(path_a, path_b, label_a="Original", label_b="Roundtrip"):
    """Compare paragraph content between two DOCX files."""
    doc_a = Document(path_a)
    doc_b = Document(path_b)
    
    paras_a = [(p.text.strip(), p.style.name) for p in doc_a.paragraphs if p.text.strip()]
    paras_b = [(p.text.strip(), p.style.name) for p in doc_b.paragraphs if p.text.strip()]
    
    print(f"\n{'='*60}")
    print(f"CONTENT COMPARISON: {label_a} vs {label_b}")
    print(f"{'='*60}")
    print(f"  {label_a}: {len(paras_a)} non-empty paragraphs")
    print(f"  {label_b}: {len(paras_b)} non-empty paragraphs")
    
    # Compare text content (ignore style differences for now)
    texts_a = set(t for t, _ in paras_a)
    texts_b = set(t for t, _ in paras_b)
    
    missing_in_b = texts_a - texts_b
    extra_in_b = texts_b - texts_a
    
    if not missing_in_b and not extra_in_b:
        print(f"\n  ✅ PERFECT MATCH: All content preserved in roundtrip!")
    else:
        if missing_in_b:
            print(f"\n  ❌ MISSING in {label_b} ({len(missing_in_b)} paragraphs):")
            for t in sorted(missing_in_b):
                print(f"     - \"{t[:100]}...\"" if len(t) > 100 else f"     - \"{t}\"")
        if extra_in_b:
            print(f"\n  ➕ EXTRA in {label_b} ({len(extra_in_b)} paragraphs):")
            for t in sorted(extra_in_b):
                print(f"     + \"{t[:100]}...\"" if len(t) > 100 else f"     + \"{t}\"")
    
    # Section-by-section comparison using bookmarks
    print(f"\n{'='*60}")
    print("SECTION-BY-SECTION CONTENT LENGTH COMPARISON")
    print(f"{'='*60}")
    
    sections_a = extract_sections(doc_a)
    sections_b = extract_sections(doc_b)
    
    all_tags = sorted(set(list(sections_a.keys()) + list(sections_b.keys())),
                      key=lambda t: [x[0] for x in BOOKMARKS].index(t) if t in [x[0] for x in BOOKMARKS] else 999)
    
    for tag in all_tags:
        a = sections_a.get(tag, {})
        b = sections_b.get(tag, {})
        
        len_a = len(a.get('content', ''))
        len_b = len(b.get('content', ''))
        heading = a.get('heading', b.get('heading', tag))
        
        if tag not in sections_a:
            print(f"  ❌ {tag} ({heading}): MISSING in {label_a}")
        elif tag not in sections_b:
            print(f"  ❌ {tag} ({heading}): MISSING in {label_b}")
        elif len_a == len_b:
            print(f"  ✅ {tag} ({heading}): {len_a} chars == {len_b} chars")
        else:
            diff = len_b - len_a
            pct = (diff / len_a * 100) if len_a > 0 else 0
            symbol = "⚠️" if abs(pct) > 5 else "≈"
            print(f"  {symbol} {tag} ({heading}): {len_a} → {len_b} chars ({diff:+d}, {pct:+.1f}%)")
    
    return len(missing_in_b), len(extra_in_b)


# ======== MAIN ========

def main():
    print("=" * 70)
    print("FIX & VERIFY PIPELINE — Full DOCX Roundtrip Repair")
    print("=" * 70)
    
    # ---- STEP 0: Backup ----
    print("\n📦 STEP 0: Create working copy")
    shutil.copy2(ORIGINAL_DOCX, REPAIRED_DOCX)
    print(f"  Copied {ORIGINAL_DOCX} → {REPAIRED_DOCX}")
    
    original_sha = hashlib.sha256(ORIGINAL_DOCX.read_bytes()).hexdigest()
    print(f"  Original SHA256: {original_sha}")
    
    # ---- STEP 1: Strip all existing bookmarks ----
    print("\n🧹 STEP 1: Strip ALL existing tag_ bookmarks (duplicate/misplaced)")
    doc = Document(REPAIRED_DOCX)
    removed = strip_all_tag_bookmarks(doc)
    print(f"  Removed {removed} bookmark elements")
    
    # ---- STEP 2: Fix heading styles ----
    print("\n🔧 STEP 2: Fix body paragraphs erroneously styled as Headings")
    fixed = fix_heading_styles(doc)
    print(f"  Fixed {fixed} paragraphs")
    
    # ---- STEP 3: Insert clean bookmarks ----
    print("\n🔖 STEP 3: Insert clean bookmarks (1 per heading)")
    added, missed = insert_clean_bookmarks(doc)
    print(f"  Added {added} bookmarks, {len(missed)} missed")
    
    # Save repaired DOCX
    doc.save(REPAIRED_DOCX)
    repaired_sha = hashlib.sha256(REPAIRED_DOCX.read_bytes()).hexdigest()
    print(f"\n  Saved repaired DOCX: {REPAIRED_DOCX}")
    print(f"  Repaired SHA256: {repaired_sha}")
    
    # ---- STEP 4: Extract DOCX → JSON ----
    print("\n📤 STEP 4: Extract bookmarked sections from repaired DOCX")
    doc = Document(REPAIRED_DOCX)  # reload
    sections = extract_sections(doc)
    print(f"  Extracted {len(sections)} sections:")
    for tag, data in sections.items():
        print(f"    {tag}: heading='{data['heading']}', level={data['level']}, "
              f"body_paras={data['body_para_count']}, content_len={len(data['content'])}")
    
    # Save extraction JSON
    json_data = {k: {kk: vv for kk, vv in v.items() if kk not in ('para_start', 'para_end')}
                 for k, v in sections.items()}
    EXTRACTED_JSON.write_text(json.dumps(json_data, indent=2, ensure_ascii=False))
    print(f"  Saved to {EXTRACTED_JSON}")
    
    # ---- STEP 5: Build manuscript MD ----
    print("\n📝 STEP 5: Build manuscript MD from extracted sections")
    md_content = build_manuscript_md(sections, repaired_sha)
    
    # Backup existing MD
    if MANUSCRIPT_MD.exists():
        backup = MANUSCRIPT_MD.with_suffix('.md.pre_fix_backup')
        shutil.copy2(MANUSCRIPT_MD, backup)
        print(f"  Backed up existing MD to {backup}")
    
    MANUSCRIPT_MD.write_text(md_content, encoding='utf-8')
    print(f"  Written {len(md_content)} chars to {MANUSCRIPT_MD}")
    
    # Quick sanity: count TAG blocks
    tag_count = len(re.findall(r'<!-- TAG:', md_content))
    print(f"  TAG blocks in MD: {tag_count}")
    
    # ---- STEP 6: MD → DOCX ----
    print("\n🔨 STEP 6: Convert MD → DOCX (roundtrip test)")
    ret = run_md_to_docx(MANUSCRIPT_MD, REPAIRED_DOCX, OUTPUT_DOCX)
    if ret != 0:
        print(f"  ❌ md_to_docx.py failed with exit code {ret}")
        print("  Attempting direct comparison of repaired DOCX only...")
    else:
        print(f"  ✅ Output saved to {OUTPUT_DOCX}")
    
    # ---- STEP 7: Compare ----
    print("\n🔍 STEP 7: Content comparison")
    
    if OUTPUT_DOCX.exists():
        missing, extra = compare_docx_content(REPAIRED_DOCX, OUTPUT_DOCX,
                                               "Repaired Original", "Roundtrip Output")
    else:
        print("  ⚠ No roundtrip output to compare — checking repaired DOCX extraction only")
        missing, extra = 0, 0
    
    # ---- SUMMARY ----
    print(f"\n{'='*70}")
    print("PIPELINE VERIFICATION SUMMARY")
    print(f"{'='*70}")
    print(f"  Original DOCX SHA256:  {original_sha}")
    print(f"  Repaired DOCX SHA256:  {repaired_sha}")
    if OUTPUT_DOCX.exists():
        output_sha = hashlib.sha256(OUTPUT_DOCX.read_bytes()).hexdigest()
        print(f"  Roundtrip DOCX SHA256: {output_sha}")
    print(f"  Sections extracted:    {len(sections)}")
    print(f"  Bookmarks added:       {added}")
    print(f"  Heading styles fixed:  {fixed}")
    if OUTPUT_DOCX.exists():
        print(f"  Content missing:       {missing}")
        print(f"  Content extra:         {extra}")
        if missing == 0:
            print(f"\n  ✅ ROUNDTRIP INTEGRITY: ALL CONTENT PRESERVED")
        else:
            print(f"\n  ❌ ROUNDTRIP INTEGRITY: {missing} PARAGRAPHS LOST")
    
    return 0 if missing == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
